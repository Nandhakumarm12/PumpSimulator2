"""
Build Paper 2 Word document using python-docx.
Title: From Manual to Model: Simulation-Enabled Safety and Security Research for Infusion Pumps
Venue: Prototypes for Humanity (Dubai)
Run from project root: python3 docs/build_paper2_docx.py
Output: docs/paper2_from_manual_to_model.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Styles ──────────────────────────────────────────────────────────────────
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.bold = True
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.italic = False
    return p

def body(text, italic=False, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = italic
    run.bold   = bold
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.italic = True
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.4)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p

def ref_entry(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent        = Inches(0.4)
    p.paragraph_format.first_line_indent  = Inches(-0.4)
    p.paragraph_format.space_after        = Pt(4)
    p.paragraph_format.line_spacing       = Pt(13)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for para in hdr.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
            for para in tr.cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE & AUTHORS
# ══════════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(12)
tr = title_p.add_run('From Manual to Model: Simulation-Enabled Safety and Security Research for Infusion Pumps')
tr.font.name = 'Times New Roman'
tr.font.size = Pt(16)
tr.bold = True

auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(4)
ar = auth_p.add_run('[Author names and affiliations — to be completed]')
ar.font.name = 'Times New Roman'
ar.font.size = Pt(11)
ar.italic = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════

h1('Abstract')
body(
    'Infusion pump safety research is structurally constrained by a physical access problem: '
    'the devices most implicated in serious medication errors are the most expensive and '
    'operationally restricted to acquire for research purposes. A single large-volume pump '
    'with drug library licensing can exceed £10,000, and comparative studies spanning multiple '
    'device families are rarely feasible outside manufacturer partnerships. This paper presents '
    'a methodology — and the open-source simulation platform it produced — for building '
    'behaviorally faithful software replicas of clinical infusion pumps directly from '
    'manufacturer Directions for Use (DFU) documents. The methodology follows four reproducible '
    'steps: DFU analysis, pure functional state machine encoding, guardrail logic implementation, '
    'and automated data pipeline construction. Three complete device simulators are described: '
    'the BD Alaris GP volumetric pump (two-tier guardrail, ten screen states, validated against '
    'the published PVSio-web formal model), the B. Braun Infusomat Space (three-tier guardrail, '
    'SpaceCom2 network module, firmware vulnerabilities CVE-2021-33885 and CVE-2021-33882), '
    'and the Graseby 3100 syringe driver (zero guardrails, the defining clinical safety risk '
    'of its era). A 58-test Behavioural Fidelity Validation suite — in which every assertion '
    'is traceable to a specific DFU section, formal model reference, or published CVE advisory '
    '— achieves 100% pass rate across all three devices. The platform generates structured '
    'training datasets for AI-based safety and security research at thousands of scenarios '
    'per minute, and enables cross-device comparative analysis that would be logistically '
    'infeasible with physical hardware. The codebase comprises approximately 18,000 lines of '
    'TypeScript and is released as open-source software.'
)
body('Keywords: infusion pump simulation, behavioural fidelity, medical device safety, '
     'guardrail systems, cybersecurity, open-source research platform', italic=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

h1('1  Introduction')

body(
    'Between 2005 and 2009, the United States Food and Drug Administration (FDA) received '
    'over 56,000 medical device reports linked to infusion pump malfunction or misuse, '
    'with 710 associated patient deaths [1]. This concentration of adverse events in a '
    'single device category reflects a fundamental tension in clinical practice: infusion '
    'pumps deliver the most dangerous medications — vasopressors, opioids, anticoagulants, '
    'concentrated electrolytes — through an interface that was designed before modern '
    'software safety engineering and has changed little since. The chevron-based number '
    'entry model used by the Alaris GP, the B. Braun Infusomat Space, and many other '
    'contemporary large-volume pumps was mathematically analysed by Thimbleby and Cairns [2], '
    'who demonstrated that ten-fold dose errors are a systematic, predictable property of '
    'this interaction design, not an occasional user failure. Cauchi et al. [3] later '
    'formalised this finding using the PVSio verification framework, producing machine-checked '
    'proofs of error conditions in the Alaris GP interface. Despite this research foundation, '
    'comparative safety studies spanning multiple device families remain rare in the literature. '
    'The reason is practical: physical access.'
)

body(
    'Infusion pumps are tightly controlled medical devices. Acquiring three or more pump '
    'models for research use requires institutional purchase, regulatory permissions, clinical '
    'governance approval, and access to trained staff for safe operation. Beyond the direct '
    'cost — which can exceed £30,000 for a three-device comparative study — physical pumps '
    'cannot be run at scale: generating 1,000 simulated clinical scenarios on a real device '
    'would require 1,000 manual programming sessions. This access barrier has concentrated '
    'infusion pump safety research in well-resourced clinical institutions and manufacturer '
    'partnerships, leaving the broader research community with limited ability to study '
    'device families they cannot afford to acquire.'
)

body(
    'Simulation offers a principled alternative, but prior computational approaches to '
    'infusion pump research have been constrained by their own limitations. Formal verification '
    'tools — PVS, Alloy, Event-B — produce rigorous proofs but require expert theorem prover '
    'skills and yield models that are not directly executable as research platforms [4]. '
    'Higher-level simulation environments used in human factors research are typically '
    'purpose-built for a single device, not designed for automated dataset generation, and '
    'not validated against manufacturer documentation with traceable evidence [5]. What has '
    'been absent is a middle layer: a device simulator that is (a) faithful enough to '
    'generate ecologically valid clinical scenarios, (b) executable without specialist '
    'infrastructure, (c) extensible across multiple device families, and (d) validated with '
    'transparent, reproducible evidence.'
)

body(
    'This paper presents exactly that platform, produced as part of a funded AI safety and '
    'security ranking program for medical device fleets. The key contribution is methodological: '
    'a four-step process for translating any publicly available DFU document into a '
    'research-grade, validated device simulator. The methodology is demonstrated across three '
    'devices selected to represent a spectrum of design complexity and safety risk — from '
    'the Alaris GP\'s formal-model-validated two-tier guardrail architecture, through the '
    'B. Braun Infusomat Space\'s three-tier system and confirmed network-layer CVEs, to the '
    'Graseby 3100\'s complete absence of dosing safeguards. Together, these three simulators '
    'enable safety and security analyses that span the full range of devices currently in '
    'active clinical use in UK hospitals.'
)

body(
    'The remainder of the paper is structured as follows. Section 2 provides background on '
    'prior formal modelling and simulation work for infusion pumps. Section 3 presents the '
    'From Manual to Model methodology. Section 4 describes the three device simulators in '
    'detail. Section 5 reports the Behavioural Fidelity Validation results. Section 6 '
    'describes the research the platform enables. Section 7 discusses limitations and '
    'future work, and Section 8 concludes.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. BACKGROUND AND RELATED WORK
# ══════════════════════════════════════════════════════════════════════════════

h1('2  Background and Related Work')

h2('2.1  Formal Modelling of Infusion Pumps')

body(
    'The PVSio-web project represents the most substantial prior work in formal pump modelling. '
    'Masci et al. [4] produced a PVS-verified formal model of the Alaris GP that defines '
    'state transitions as machine-checked specifications and provides a web-based prototype '
    'interface for user testing. The PVSio-web model covers the full rate entry workflow: '
    'language selection, drug selection, rate entry via chevrons, guardrail enforcement, '
    'infusion start, hold, and alarm handling. It has been used to identify formal properties '
    'of number entry errors and to study the gap between specification and observed user '
    'behaviour [3]. The Alaris GP simulator described in this paper was developed independently '
    'and then cross-checked against the PVSio-web state transition definitions; 10 of 25 '
    'Alaris GP validation tests explicitly verify alignment with this formal model.'
)

body(
    'The critical limitation of PVSio-web for our purpose is scope: it models a single '
    'device, does not extend to data generation at scale, and was not designed for '
    'multi-device comparative research. Masci et al. [4] describe the model as a platform '
    'for user study design and interface evaluation, not for generating structured training '
    'datasets. Its value to our work is as a secondary validation authority for Alaris GP '
    'behaviour — a role it fulfils with rigour that manufacturer documentation alone cannot '
    'provide.'
)

h2('2.2  Guardrail Research and DERS Effectiveness')

body(
    'Drug library guardrails — tiered rate limits that warn or block when a programmed '
    'rate exceeds clinical thresholds — are the primary software safety mechanism in modern '
    'large-volume pumps. Research on guardrail effectiveness has documented a consistent '
    'and clinically significant finding: guardrail override rates in hospital deployments '
    'typically range from 35% to 95% depending on drug and care area [6]. Studies of '
    'smart pump deployments find that DERS effectiveness depends critically on drug library '
    'currency, limit calibration, and clinical workflow integration [7]. When any of these '
    'conditions degrade — outdated libraries, broadly set limits, alert fatigue — the '
    'guardrail system provides limited protection despite its presence. Studying these '
    'dynamics with physical devices requires controlled clinical studies with ethical '
    'approvals, patient risk, and operational disruption. Simulation-based research at '
    'scale, with explicit logging of every guardrail interaction and override event, '
    'is an alternative that no prior platform has fully realised.'
)

h2('2.3  Medical Device Cybersecurity')

body(
    'The cybersecurity vulnerability profile of networked infusion pumps became a major '
    'research and regulatory concern following the publication of CISA advisory '
    'ICSMA-21-294-01 [8], which disclosed multiple critical vulnerabilities in the '
    'B. Braun Infusomat Space system. CVE-2021-33885 (CVSS 9.0) enables unauthenticated '
    'remote modification of infusion parameters via the SpaceStation network protocol. '
    'CVE-2021-33882 (CVSS 8.2) exploits missing authentication in the SpaceCom2 network '
    'module. These vulnerabilities have not been fully remediated: B. Braun\'s firmware '
    'update protocol for the Infusomat Space does not incorporate cryptographic signature '
    'verification in any published firmware version [8]. A simulation platform that models '
    'network state, firmware integrity flags, and connectivity alongside clinical programming '
    'behaviour — as the platform described in this paper does — enables combined safety-security '
    'risk analysis that no prior simulator has supported.'
)

h2('2.4  Legacy Device Safety: The Graseby 3100')

body(
    'The Graseby 3100 syringe driver is not a networked device and has no documented '
    'cybersecurity vulnerabilities. Its clinical risk arises from a different source: '
    'the complete absence of dosing safeguards. The Graseby 3100 was widely deployed in '
    'UK palliative care and ICU settings from the late 1980s through the 2000s for '
    'continuous subcutaneous infusions of morphine, diamorphine, and other opioids. '
    'The National Patient Safety Agency [9] documented multiple fatal overdose incidents '
    'attributable to Graseby syringe driver programming errors and issued a patient safety '
    'alert in 2010. The device\'s rate range of 0.1–199.9 ml/h, entirely unconstrained '
    'by guardrails or drug library enforcement, makes any programming error directly '
    'consequential. Including the Graseby 3100 in the simulator platform provides a '
    'baseline against which guardrail effectiveness can be quantified by comparing '
    'clinical scenarios run on all three devices.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
#  3. FROM MANUAL TO MODEL: THE METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════

h1('3  From Manual to Model: The Methodology')

body(
    'The methodology for building each simulator follows four sequential steps. The steps '
    'are independent of any specific device — they apply whenever a manufacturer DFU or '
    'equivalent operations manual is publicly available. Figure 1 illustrates the '
    'pipeline from source documentation to research output.'
)

h2('3.1  Step 1: DFU Analysis')

body(
    'The DFU is the legally binding document describing how the device operates. It specifies '
    'every operator-visible behaviour: what each button does, what each screen displays, '
    'how the device responds to alarm conditions, and what the factory-default parameter '
    'values are. The analysis extracts four classes of information.'
)

body(
    'First, the device state inventory: every named condition the device can be in, '
    'corresponding to a screen the operator sees. For the Alaris GP, the DFU describes '
    'eleven distinct screens from LANGUAGE_SELECT through to PRESSURE_VIEW. For the Graseby '
    '3100, the equivalent manual identifies five states, reflecting the device\'s much '
    'simpler interaction model. Each state name becomes a member of a TypeScript discriminated '
    'union type, making it impossible to represent an undocumented state.'
)

body(
    'Second, the transition table: for each state, what operator action causes a transition '
    'to which other state, and under what condition. These transitions are the core of the '
    'state machine. Where the DFU specifies a guard condition — such as "RUN pressed: '
    'if rate within limits, transition to RUNNING; if above soft limit, transition to '
    'GUARDRAIL_WARNING" — the guard is encoded as a pure TypeScript predicate that can '
    'be tested independently.'
)

body(
    'Third, factory defaults: every numerical parameter that the device ships with as '
    'standard. These are stored in a constants file as a single typed object '
    '(FACTORY_DEFAULTS for the Alaris GP, BRAUN_DEFAULTS for the B. Braun) from which '
    'all other code reads. No numerical literal appears elsewhere in the codebase; '
    'changing a factory default requires a single edit to a single constant.'
)

body(
    'Fourth, ambiguity resolution: where the DFU leaves behavior unspecified — such as '
    'the exact timing parameters for chevron hold-to-accelerate behavior, which the Alaris '
    'GP DFU describes qualitatively — the PVSio-web formal model [4] is used as a secondary '
    'reference. Where neither source resolves an ambiguity, the interpretation most '
    'consistent with clinical safety is adopted and documented in the codebase.'
)

h2('3.2  Step 2: State Machine Encoding')

body(
    'Each device\'s behavior is encoded as a pure functional state machine: a module of '
    'TypeScript functions in which each function accepts the current device state and '
    'returns both the new state and a log of events generated by the transition. No '
    'mutation occurs: every state object is treated as immutable, and every transition '
    'produces a new state object. This design makes the state machine deterministic, '
    'side-effect-free, and independently testable without a browser or React environment.'
)

body('The function signature pattern is uniform across all three simulators:')

code_block(
    'function pressRun(\n'
    '  state: PumpState,\n'
    '  timestamp: number\n'
    '): ActionResult {\n'
    '  // state: current device state (immutable)\n'
    '  // timestamp: ms since session start (for log entries)\n'
    '  // returns: { state: PumpState, logEntries: SessionLogEntry[] }\n'
    '}'
)

body(
    'The state machine module has a strict architectural constraint: zero React imports '
    'are permitted. This is enforced by the TypeScript compiler configuration and verified '
    'by the test suite, which runs in a Node.js environment with no browser globals. '
    'The React UI layer calls state machine functions through hooks; it never implements '
    'clinical logic directly. This separation makes the clinical logic auditable without '
    'understanding the UI implementation.'
)

body(
    'Every log entry is frozen (Object.freeze) at creation to enforce the session log '
    'immutability requirement. Log entries are never mutated after creation; they can '
    'only be appended to the session log array. This ensures that the dataset generated '
    'from a session accurately reflects the sequence of events as they occurred, without '
    'any possibility of retrospective modification.'
)

h2('3.3  Step 3: Guardrail Logic Implementation')

body(
    'Guardrail logic is implemented as a separate module that receives a rate in the '
    'drug\'s native units and a drug object from the library, and returns a guardrail '
    'status. The separation from the state machine allows guardrail logic to be tested '
    'independently, replaced for different device families, and audited against clinical '
    'sources without touching the state machine code.'
)

body(
    'The guardrail interface is consistent across devices despite their different tier '
    'structures. The Alaris GP checkGuardrail() function returns one of three statuses: '
    '"ok", "warning" (soft limit exceeded), or "blocked" (hard limit exceeded). The '
    'B. Braun checkBraunGuardrail() function returns four statuses, adding "advisory" '
    'for rates within 20% of the soft limit boundary. The Graseby 3100 has no guardrail '
    'function at all — pressStart() in the Graseby state machine calls no guardrail check '
    'and transitions directly to RUNNING for any non-zero rate.'
)

body(
    'Drug library data is stored as typed objects with six threshold fields per drug: '
    'softMin, softMax, hardMin, hardMax, defaultRate, and weightBased. Threshold values '
    'are sourced from clinical practice guidelines and are not adjustable at runtime. '
    'Changing a guardrail threshold requires a code modification with a traceable '
    'justification — exactly the governance behaviour required for a real device\'s '
    'drug library update.'
)

h2('3.4  Step 4: Automated Data Pipeline')

body(
    'Each state machine is connected to a scenario generator that drives the simulator '
    'programmatically through clinical scenarios without human interaction. A scenario '
    'specifies: (a) a device profile describing system-level state (maintenance age, '
    'battery level, firmware version, network connectivity, drug library currency); '
    '(b) a clinical sequence (drug selection, target rate, number of chevron presses, '
    'guardrail response if triggered); and (c) a session outcome (whether the programmed '
    'rate was correct, whether the nurse overrode a guardrail, whether VTBI was set).'
)

body(
    'The scenario generator runs each scenario through the state machine, collecting the '
    'resulting session log. A feature extractor then processes the log into a 48-field '
    'training record covering interaction metrics, configuration choices, and system state. '
    'A rule engine applies 26 evidence-based risk rules — 21 for the Alaris GP and B. Braun, '
    'with 5 additional device-specific rules for the B. Braun\'s network and firmware '
    'properties — to assign a risk label and continuous risk score. The pipeline generates '
    '1,000 training records across all three devices in under 30 seconds on consumer '
    'hardware.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
#  4. DEVICE SIMULATORS
# ══════════════════════════════════════════════════════════════════════════════

h1('4  Device Simulators')

h2('4.1  BD Alaris GP Volumetric Infusion Pump')

body(
    'The BD Alaris GP (CareFusion, now BD) is one of the most widely deployed large-volume '
    'infusion pumps in NHS hospitals and the primary device used in the CHI-MED formal '
    'modelling research [3,4]. The simulator implements all eleven documented screen states, '
    'ten drugs from the clinical library, two-tier guardrail enforcement, VTBI programming, '
    'and four alarm types (OCCLUSION, AIR_IN_LINE, INFUSION_COMPLETE, BATTERY_LOW) with '
    'the priority ordering specified in the DFU: OCCLUSION takes precedence over AIR_IN_LINE, '
    'which takes precedence over BATTERY_LOW, which takes precedence over INFUSION_COMPLETE. '
    'The state machine comprises 503 lines of TypeScript implementing 22 pure functions.'
)

body(
    'The guardrail system is sourced from the Guardrails Edition of the DFU (BDDF00535 '
    'Issue 4), which describes the soft and hard limit behavior in detail. A rate above '
    'the drug\'s softMax triggers a GUARDRAIL_WARNING screen requiring the nurse to choose '
    'OVERRIDE or RE-ENTER; a rate above hardMax triggers GUARDRAIL_BLOCKED, from which '
    'only RE-ENTER is available. OVERRIDE events are logged as a separate event type '
    '(guardrail_override) and are never silently captured within another event — this '
    'logging requirement is the most safety-critical invariant in the codebase, since '
    'guardrail override behavior is the primary outcome variable in DERS effectiveness '
    'research [6].'
)

body(
    'The chevron rate entry system implements hold-to-accelerate: a 500ms initial hold '
    'delay before repeat firing at 80ms intervals, matching the DFU specification [10]. '
    'Double chevrons apply a step of 10 ml/h (STEP_LARGE); single chevrons apply 1 ml/h '
    '(STEP_SMALL). Rate boundaries are absolute: at RATE_MAX (1200 ml/h) or RATE_MIN '
    '(0.1 ml/h), the rate is clamped and a boundary_hit event is logged. Direction '
    'reversals during entry are detected by comparing the sign of the current chevron '
    'press with the last rate_adjust event, and logged as correction events — the metric '
    'used by Cauchi et al. [3] to characterise number entry difficulty.'
)

body(
    'The Alaris GP simulator benefits from a secondary validation source not available '
    'for the other two devices: the PVSio-web formal model [4]. This model defines the '
    'Alaris GP state machine as a PVS specification and has been peer-reviewed and '
    'published independently of the manufacturer. Ten of the 25 Alaris GP behavioural '
    'fidelity tests verify alignment with the PVSio-web specification for the core '
    'programming workflow — language select, drug select, rate entry, guardrail handling, '
    'run, hold, and alarm.'
)

h2('4.2  B. Braun Infusomat Space Large-Volume Pump')

body(
    'The B. Braun Infusomat Space introduces four design differences from the Alaris GP '
    'that are directly relevant to safety and security research, each faithfully captured '
    'in the simulator.'
)

body(
    'The three-tier guardrail system is the most clinically significant difference. Below '
    'the soft limit boundary, the Infusomat Space adds an advisory zone covering the region '
    'within 20% of the boundary (ADVISORY_ZONE_FRACTION = 0.20). A rate in this zone '
    'triggers a GUARDRAIL_ADVISORY screen requiring nurse acknowledgement — a single '
    'keypress that confirms awareness but does not require clinical justification. This '
    'behavioral distinction matters for research: advisory acknowledgement is qualitatively '
    'different from override. Overriding a soft limit requires the nurse to accept clinical '
    'responsibility for exceeding the recommended range; acknowledging an advisory merely '
    'confirms awareness of proximity to that range. The three-tier system thus produces a '
    'richer behavioral signal — three event types rather than two — for the same clinical '
    'programming sequence.'
)

body(
    'The SpaceCom2 network module is modelled as a binary flag (spacecom2Connected) in '
    'the device state. Connecting or disconnecting the module generates a dedicated log '
    'entry (spacecom2_connected / spacecom2_disconnected). When connected, the device is '
    'exposed to CVE-2021-33882 [8] — missing authentication for network commands — which '
    'the simulator flags as a risk factor in the generated training data. The connectivity '
    'state also affects drug library update capability: a disconnected pump cannot receive '
    'library updates from SpaceStation, making drug_library_age_days a compound risk factor '
    'that interacts with connectivity state.'
)

body(
    'Firmware unsigned status (CVE-2021-33885, CVSS 9.0 [8]) is encoded as a permanent '
    'device property: firmwareSigned = false in both the constants file and the initial '
    'state. This is not a runtime configuration; it reflects the confirmed finding that '
    'no firmware version in the Infusomat Space series implements cryptographic signature '
    'verification. The simulator correctly treats this as an immutable design property '
    'rather than a session variable. The B. Braun state machine comprises 1,195 lines '
    'of TypeScript — larger than the Alaris GP (503 lines) due to the additional guardrail '
    'tier, SpaceCom2 module functions, and firmware recall toggle capability.'
)

body(
    'Bolus volume is capped at 2.0 ml (BRAUN_DEFAULTS.BOLUS_MAX_ML), compared to 5.0 ml '
    'on the Alaris GP. The device starts at a STARTUP screen (self-test) rather than '
    'LANGUAGE_SELECT. Rate is edited directly in the state rate field during entry, '
    'without a separate rateBuffer — reflecting the B. Braun\'s single-field rate entry '
    'model as documented in the IFU.'
)

h2('4.3  Graseby 3100 Syringe Driver')

body(
    'The Graseby 3100 is represented in the simulator as a deliberately minimal device: '
    'five screen states (BOOT, RATE_ENTRY, RUNNING, ON_HOLD, ALARM), no drug library, '
    'no guardrail module, no VTBI, and no bolus mode. The state machine comprises 323 '
    'lines of TypeScript — the shortest of the three, reflecting the device\'s simplified '
    'interaction model. The complete absence of guardrail logic is architecturally enforced: '
    'the pressStart() function contains no call to any guardrail check and transitions '
    'directly to RUNNING for any rate above zero within the device\'s range. This is not '
    'an implementation gap — it is a faithful representation of how the device operates.'
)

body(
    'The volume constraint is syringe capacity rather than VTBI. Three standard syringe '
    'sizes are supported (20 ml, 30 ml, 50 ml), matching the BD Plastipak and Monoject '
    'syringes listed in the Graseby 3100 Operators Manual [11]. When accumulated volume '
    'reaches the selected capacity, a SYRINGE_EMPTY alarm fires and infusion stops. '
    'This models the real device\'s mechanical stop mechanism, which prevents further '
    'delivery when the syringe plunger reaches the barrel limit but does not alert '
    'the nurse until the device detects that delivery has halted.'
)

body(
    'The rate range (0.1–199.9 ml/h, GRASEBY_DEFAULTS.RATE_MAX = 199.9) is lower '
    'than the large-volume pumps (1200 ml/h) because the Graseby 3100 is a syringe '
    'driver, not a peristaltic pump. The NPSA alert [9] that documented fatal incidents '
    'with the Graseby 3100 found that the most common programming error was a tenfold '
    'rate error — for example, entering 20 ml/h instead of 2 ml/h for a diamorphine '
    'infusion. The simulator captures this scenario through the same correction detection '
    'and boundary hit logging used by the Alaris GP, enabling direct comparison of '
    'programming error patterns across devices for the same target rate.', space_after=10
)

table(
    headers=['Property', 'Alaris GP', 'B. Braun Infusomat', 'Graseby 3100'],
    rows=[
        ['Guardrail tiers',        '2 (warning + blocked)', '3 (advisory + warning + blocked)', '0 (none)'],
        ['Drug library',           'Yes (10 drugs, sim)',   'Yes (1,500 entries, 10 in sim)', 'No'],
        ['VTBI',                   'Yes',                   'Yes',                             'No (syringe capacity)'],
        ['Bolus maximum',          '5.0 ml',                '2.0 ml',                          'N/A'],
        ['Rate maximum',           '1200 ml/h',             '1200 ml/h',                       '199.9 ml/h'],
        ['Network module',         'None',                  'SpaceCom2 (WiFi)',                 'None'],
        ['Firmware CVEs',          'None documented*',      'CVE-2021-33885, CVE-2021-33882',  'None documented'],
        ['Initial screen',         'LANGUAGE_SELECT',       'STARTUP',                         'RATE_ENTRY'],
        ['State machine (LOC)',    '503',                   '1,195',                           '323'],
        ['Formal model available', 'Yes (PVSio-web [4])',   'No',                              'No'],
    ],
    col_widths=[1.65, 1.35, 1.85, 1.45]
)
caption('Table 1: Cross-device simulator comparison. *BD Alaris has published CVEs for the '
        'Alaris system family (CISA 2023) but these are outside the scope of this paper. '
        'LOC = lines of TypeScript in the state machine module only.')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  5. BEHAVIOURAL FIDELITY VALIDATION
# ══════════════════════════════════════════════════════════════════════════════

h1('5  Behavioural Fidelity Validation')

h2('5.1  Validation Framework')

body(
    'A simulator that does not behave like the real device generates invalid data. '
    'Validating behavioural fidelity without access to physical hardware requires a '
    'structured argument from documentation rather than device-to-device comparison. '
    'We operationalise this as a Behavioural Fidelity Validation (BFV) suite in which '
    'every test assertion is explicitly anchored to one of three types of primary source: '
    'a specific DFU section, a published formal model specification, or a manufacturer-confirmed '
    'cybersecurity advisory. Tests are implemented as pure TypeScript functions integrated '
    'with the Vitest test runner, so they run as part of the standard npm test command '
    'and appear in CI results alongside unit tests. A separate report generator produces '
    'a structured JSON artefact (data/validation/report.json) containing the full test '
    'results with DFU references, expected values, actual simulator outputs, and '
    'PVSio-web alignment flags — the primary evidence document for this section of the paper.'
)

body(
    'Each BFV test case has five components: a unique identifier following the pattern '
    'BFV-{device}-{number} (e.g., BFV-AG-009); a description; a primary source reference; '
    'a human-readable expected outcome; and an executable assertion that evaluates the '
    'simulator\'s actual output. An optional sixth component, pvsioweb, flags Alaris GP '
    'tests that are also validated against the PVSio-web formal model as a secondary '
    'authority. Tests are grouped into four suites: Alaris GP (DFU BD 1000DF00152 and '
    'BDDF00535), B. Braun Infusomat Space (IFU 686N-GB and CISA ICSMA-21-294-01), '
    'Graseby 3100 (Operators Manual 2002 and NPSA Alert 2010), and cross-device '
    'comparison.'
)

h2('5.2  Test Coverage and Results')

body('Table 2 presents the full BFV test coverage and results.')

table(
    headers=['Device Suite', 'Tests', 'Passed', 'Failed', 'Pass Rate', 'PVSio-web Aligned'],
    rows=[
        ['Alaris GP (BD 1000DF00152 / BDDF00535)',      '25', '25', '0', '100%', '10'],
        ['B. Braun Infusomat Space (IFU 686N-GB / CISA)', '18', '18', '0', '100%', 'N/A'],
        ['Graseby 3100 (Operators Manual / NPSA 2010)', '10', '10', '0', '100%', 'N/A'],
        ['Cross-Device Comparison',                       '5',  '5',  '0', '100%', 'N/A'],
        ['Total',                                         '58', '58', '0', '100%', '10'],
    ],
    col_widths=[2.8, 0.6, 0.65, 0.6, 0.85, 1.2]
)
caption('Table 2: Behavioural Fidelity Validation results. PVSio-web Aligned = number of '
        'tests additionally verified against the published formal model.')

body(
    'The Alaris GP suite covers: initial state (LANGUAGE_SELECT), all valid transitions '
    'in the core programming workflow, chevron step magnitudes (STEP_LARGE=10, STEP_SMALL=1 '
    'matching DFU Table 1), boundary clamping at RATE_MIN (0.1 ml/h) and RATE_MAX '
    '(1200 ml/h), two-tier guardrail activation (soft warning and hard block), guardrail '
    'override logging (the most safety-critical logging invariant), direction-reversal '
    'correction detection, VTBI programming and confirmation, INFUSION_COMPLETE alarm '
    'when volumeInfused reaches VTBI, KVO rate activation at 1.0 ml/h (FACTORY_DEFAULTS.KVO_RATE), '
    'alarm priority ordering (OCCLUSION > AIR_IN_LINE), KVO-active alarm silence '
    'behavior, and MANUAL mode guardrail bypass.'
)

body(
    'The B. Braun suite adds tests specific to the three-tier system: advisory zone '
    'boundary precision (rate 8.0 triggers RUNNING, rate 8.1 triggers GUARDRAIL_ADVISORY '
    'for morphine with softMax=10 and ADVISORY_ZONE_FRACTION=0.20), advisory acknowledgement '
    'versus explicit override (different event types), bolus clamping at 2.0 ml, '
    'SpaceCom2 connect/disconnect event logging, firmware unsigned flag as a permanent '
    'property (firmwareSigned=false in both BRAUN_DEFAULTS and initial state), and '
    'direct rate editing without a rateBuffer field.'
)

body(
    'The Graseby suite focuses on the structural and behavioral contrasts with the '
    'guardrailed devices: immediate RATE_ENTRY on power-on (no language or startup screen), '
    'zero-guardrail acceptance of any valid rate (pressStart at 150 ml/h transitions '
    'directly to RUNNING), rejection of rate=0, RATE_MAX=199.9 ml/h confirmed in both '
    'constants and boundary behavior, SYRINGE_EMPTY alarm, and structural verification '
    'that the state type contains no drug library, VTBI, or bolus fields.'
)

body(
    'The five cross-device tests verify the comparative claims that are central to the '
    'research enabled by the platform: guardrail tier counts (Alaris=2, B.Braun=3, '
    'Graseby=0), bolus maxima (Alaris=5ml, B.Braun=2ml, Graseby N/A), shared chevron '
    'step sizes across all three devices (STEP_LARGE=10, STEP_SMALL=1), rate maxima '
    '(Alaris=B.Braun=1200 ml/h, Graseby=199.9 ml/h), and initial screens '
    '(LANGUAGE_SELECT / STARTUP / RATE_ENTRY).'
)

h2('5.3  Scope and Limitations')

body(
    'The BFV suite demonstrates that the simulators match their documentary specifications. '
    'It does not demonstrate that the simulators match the physical hardware, and this '
    'distinction is important for interpreting research findings based on simulated data. '
    'Three categories of discrepancy are possible and are acknowledged explicitly.'
)

body(
    'First, documentation incompleteness. Manufacturer DFU documents are occasionally '
    'ambiguous or silent on edge cases. The Alaris GP DFU, for example, describes '
    'chevron acceleration qualitatively ("faster" and "slower") without specifying the '
    '500ms / 80ms timing parameters, which are sourced from the PVSio-web model [4]. '
    'Where our interpretation differs from the hardware\'s actual behaviour, the simulator '
    'will be wrong in a way the documentation-grounded tests cannot detect.'
)

body(
    'Second, firmware variation. The DFU documents referenced are the most recent publicly '
    'available editions. Earlier firmware versions may have had different guardrail thresholds, '
    'alarm priorities, or screen flows. Research conclusions should note the firmware '
    'version(s) on which the relevant DFU is based.'
)

body(
    'Third, physical interaction context. The simulator models the logical state machine, '
    'not the physical device experience. Haptic feedback, display brightness, alarm volume, '
    'and reaction time under clinical stress conditions are not captured. Studies of alarm '
    'fatigue, physical error recovery, or nurse workflow integration require physical devices '
    'in realistic clinical environments.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
#  6. RESEARCH ENABLED
# ══════════════════════════════════════════════════════════════════════════════

h1('6  Research Enabled by the Platform')

h2('6.1  AI Safety and Security Dataset Generation')

body(
    'The primary research application driving this platform is an AI-powered safety and '
    'security ranking system for infusion pump fleets [12]. Each simulator is connected '
    'to a programmatic scenario generator that drives the state machine through clinical '
    'programming sequences without human interaction. Four scenario profiles — ideal '
    '(well-maintained, connected, current library), neglected (overdue maintenance, '
    'outdated library, disconnected), cyber-risk (known firmware CVE, connected but '
    'vulnerable), and emergency (MANUAL mode only, no drug library) — define the ranges '
    'of device state parameters. For each scenario, the generator records the complete '
    'session log and extracts a 48-field feature vector covering interaction metrics '
    '(entry time, correction count, boundary hits, error magnitude), configuration choices '
    '(drug library used, guardrail override, VTBI set, bolus delivered), and system state '
    '(firmware version risk, maintenance age, battery level, network connectivity, drug '
    'library age). A rule engine applies 26 evidence-based risk rules to assign a risk '
    'label (low / medium / high) and continuous risk score (0.0–1.0). The full pipeline '
    'generates 1,000 labeled records across all three devices in under 30 seconds.'
)

h2('6.2  Cross-Device Comparative Analysis')

body(
    'Because all three simulators implement the same event log schema and feature extraction '
    'pipeline, their outputs are directly comparable. A researcher can apply the same '
    'clinical scenario — for example, programming morphine at a rate that exceeds the soft '
    'limit and observing whether the nurse overrides — across all three devices and compare '
    'the resulting session logs, feature vectors, and risk scores without any data '
    'transformation. The structural differences revealed by this comparison are themselves '
    'a research finding: the Graseby 3100 produces no guardrail-related events for the '
    'same scenario that produces GUARDRAIL_WARNING and a potential guardrail_override on '
    'the Alaris GP, and GUARDRAIL_ADVISORY followed by GUARDRAIL_WARNING on the '
    'B. Braun — three qualitatively different behavioral signatures for an identical '
    'clinical programming sequence.'
)

h2('6.3  Guardrail Override Behavior Analysis')

body(
    'Guardrail override events are captured as first-class log entries with full context: '
    'timestamp, current rate, drug name, guardrail tier, and the choice made (override, '
    're-enter, or advisory acknowledgement). This enables analysis of override patterns '
    'that has been difficult to study with physical devices because clinical observation '
    'studies require ethical approval, observer presence, and access to clinical staff '
    'during actual programming sessions. The simulator generates override scenarios '
    'programmatically at any desired scale, enabling statistical analysis of override '
    'frequency as a function of drug type, rate magnitude, device type, and guardrail '
    'tier — research questions that the existing literature has addressed only partially '
    'due to access constraints [6,7].'
)

h2('6.4  Security Research Without Device Access')

body(
    'The B. Braun simulator models CVE-2021-33885 and CVE-2021-33882 as runtime-observable '
    'device properties (firmwareSigned, spacecom2Connected) that appear in training records. '
    'This enables research into how cybersecurity risk interacts with clinical programming '
    'behavior in composite safety scoring — research that would require physical access to '
    'B. Braun hardware and potentially a controlled network environment to conduct on real '
    'devices. The simulator allows this interaction to be studied purely in software, '
    'generating training data that explicitly represents the compound risk of a networked '
    'device with unsigned firmware being programmed in MANUAL mode outside its guardrail '
    'limits.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
#  7. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════

h1('7  Discussion')

h2('7.1  The Case for Open-Source Device Simulation')

body(
    'Physical infusion pump access creates a structural barrier that concentrates medical '
    'device safety research in well-funded clinical institutions. The methodology described '
    'in this paper — DFU to state machine — is available to any researcher with access to '
    'a device manual and basic TypeScript skills. The DFU documents for the Alaris GP, '
    'B. Braun Infusomat Space, and Graseby 3100 are all publicly available through '
    'hospital equipment archives and equipment servicing resources. The tools required '
    '(TypeScript, Node.js, Vitest) are free. The resulting platform is open-source.'
)

body(
    'This democratisation argument must be tempered by the validation limitation identified '
    'in Section 5.3: a simulator built from documentation cannot guarantee behavioral '
    'equivalence with hardware. Researchers using simulation-generated data to draw '
    'conclusions about real clinical behaviour should treat the platform as a hypothesis '
    'generator: simulation findings identify patterns and thresholds worth investigating '
    'with real device studies, rather than replacing such studies.'
)

h2('7.2  Extending the Methodology to Other Devices')

body(
    'The four-step methodology is device-agnostic. The main variables affecting application '
    'effort are documentation quality and formal model availability. Modern devices with '
    'detailed DFUs and published formal models — such as the Alaris GP — yield the most '
    'confident simulators and the strongest validation claims. Legacy devices with sparse '
    'documentation — such as the Graseby 3100, whose 2002 manual contains minimal detail '
    'on alarm sequencing and timing — require more interpretation and produce less certain '
    'simulators. Devices with no publicly available DFU — increasingly common as '
    'manufacturers move to proprietary service documentation — cannot be simulated by '
    'this method without documentation access.'
)

h2('7.3  Synthetic vs. Real Interaction Data')

body(
    'All interaction features in the training dataset are synthetically generated by the '
    'scenario generator, not recorded from human participants programming the simulators. '
    'The implications of this are significant and should be stated explicitly. The '
    'behavioural metrics in Layer 1 — correction count, entry time, boundary hit count — '
    'are assigned by the scenario profile rather than emerging from observed human '
    'behaviour. This means the dataset is valid for demonstrating that the scoring '
    'framework correctly classifies known risk patterns, but it cannot reveal whether '
    'real nurses produce the correction counts, entry times, and boundary hit frequencies '
    'that the rules assume. A prospective user study using the simulator as a research '
    'platform — recording real interaction metrics from participants programming clinical '
    'scenarios — is a critical next step for empirical validation of the framework.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
#  8. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

h1('8  Conclusion')

body(
    'This paper has presented a methodology and open-source platform for building '
    'behaviorally faithful infusion pump simulators from manufacturer documentation, '
    'demonstrated across three devices spanning the range of clinical complexity: '
    'the BD Alaris GP, the B. Braun Infusomat Space, and the Graseby 3100 syringe driver. '
    'The four-step methodology — DFU analysis, pure functional state machine encoding, '
    'guardrail logic implementation, and automated data pipeline construction — is '
    'reproducible, device-agnostic, and requires no physical device access.'
)

body(
    'The 58-test Behavioural Fidelity Validation suite, achieving 100% pass rate across '
    'all three devices with every assertion traceable to a primary source, provides '
    'evidence of simulator correctness to a standard not typically reported in related '
    'work. The platform enables AI safety and security research at a scale and across '
    'a device range that would be logistically infeasible with physical hardware, '
    'and provides the data generation infrastructure for an ongoing AI-powered infusion '
    'pump fleet ranking programme.'
)

body(
    'The primary contribution of this work is methodological: the demonstration that '
    'research-grade, multi-device infusion pump simulation can be produced from '
    'publicly available documentation by a small research team, and validated with '
    'transparent, reproducible evidence. This makes infusion pump safety and security '
    'research accessible to groups that would otherwise be excluded by the cost and '
    'availability of physical devices — a non-trivial widening of who can meaningfully '
    'contribute to an important clinical problem.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

h1('References')

references = [
    '[1]  U.S. Food and Drug Administration (2010) Infusion Pump Improvement Initiative. '
    'White Paper, Center for Devices and Radiological Health, Silver Spring, MD.',

    '[2]  Thimbleby, H. and Cairns, P. (2010) \'Reducing number entry errors: solving a '
    'widespread, serious problem\', Journal of the Royal Society Interface, 7(51), '
    'pp. 1429–1439. doi:10.1098/rsif.2010.0112',

    '[3]  Cauchi, A., Curzon, P., Eslambolchilar, P., Gimblett, A., Huang, H., Lee, P., '
    'Li, Y., Masci, P., Oladimeji, P., Ruksenas, R. and Thimbleby, H. (2011) \'Towards '
    'dependable number entry for medical devices\', in Proc. EICS4Med Workshop at ACM CHI. '
    'CHI-MED Project. Available: www.chi-med.ac.uk',

    '[4]  Masci, P., Ruksenas, R., Oladimeji, P., Cauchi, A., Gimblett, A., Li, Y., '
    'Curzon, P. and Thimbleby, H. (2011) \'On formalising interactive number entry on '
    'infusion pumps\', in Proc. Fourth International Workshop on Formal Methods for '
    'Interactive Systems (FMIS 2011). Springer, pp. 81–98. PVSio-web demo: '
    'http://www.pvsioweb.org/demos/AlarisGP',

    '[5]  Rajkomar, A., Blandford, A. and Mayer, A. (2012) \'Physiological closed-loop '
    'simulation for infusion pump interface design verification\', Anaesthesia, 67(1), '
    'pp. 58–66. doi:10.1111/j.1365-2044.2011.06955.x',

    '[6]  Cassano-Piché, A.L., Fan, M., Sabovitch, J., Elke, M. and Easty, A.C. (2012) '
    '\'Multiple intravenous infusions phase 1b: Practice and training scan\', Ontario Health '
    'Technology Assessment Series, 12(16), pp. 1–132.',

    '[7]  Ohashi, K., Dalleur, O., Dykes, P.C. and Bates, D.W. (2014) \'Benefits and risks '
    'of using smart pumps to reduce medication error rates: a systematic review\', Drug '
    'Safety, 37(12), pp. 1011–1020. doi:10.1007/s40264-014-0232-7',

    '[8]  Cybersecurity and Infrastructure Security Agency (2021) ICS Medical Advisory '
    'ICSMA-21-294-01: B. Braun Infusomat Space Large Volume Pump and SpaceCom2. '
    'Washington, DC: US Department of Homeland Security.',

    '[9]  National Patient Safety Agency (2010) Safer use of syringe drivers in palliative '
    'care. NPSA Patient Safety Alert NPSA/2010/RRR019, London.',

    '[10] BD Medical (2017) Alaris GP Volumetric Infusion Pump — Directions For Use, '
    'Document 1000DF00152 Issue 1. Becton Dickinson.',

    '[11] Graseby Medical Ltd (2002) Graseby 3100 Syringe Driver — Operators Manual. '
    'Ardus Medical.',

    '[12] Samy, N. (2025) \'An AI-Empowered Safety and Security Ranking System for Infusion '
    'Pump Medical Devices\', QR Seed Pilot Study, University [Institution].',

    '[13] B. Braun Medical (2024) Infusomat Space Large Volume Pump — Instructions for '
    'Use, Document 686N-GB. B. Braun Melsungen AG.',
]

for ref in references:
    ref_entry(ref)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════

out = 'docs/paper2_from_manual_to_model.docx'
doc.save(out)
print(f'Saved: {out}')
