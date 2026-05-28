"""
Generate the PumpSimulator2 project documentation as a Word (.docx) file.
Run: python3 docs/generate_docs.py
Output: docs/PumpSimulator2_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x3a, 0x6a)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2a, 0x5a, 0x9a)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x2a, 0x6a, 0x6a)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shading)
    return p

def table_header(table, headers, header_color='1E3A6A'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), header_color)
        cell._tc.get_or_add_tcPr().append(shading)

def add_row(table, values, shade=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        if shade:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'EEF4FF')
            cell._tc.get_or_add_tcPr().append(shading)
    return row

def hr():
    doc.add_paragraph('─' * 90)

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run('\n\n\n')

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('PUMP SIMULATOR 2')
run.font.size   = Pt(28)
run.font.bold   = True
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6a)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('AI-Empowered Safety & Security Ranking System for Infusion Pump Medical Devices')
run.font.size   = Pt(14)
run.font.color.rgb = RGBColor(0x2a, 0x5a, 0x9a)

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Technical Documentation — Architecture, Design Flows, Features & Manual References')
run.font.size   = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'Generated: {datetime.date.today().strftime("%d %B %Y")}')
run.font.size   = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

funding_p = doc.add_paragraph()
funding_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = funding_p.add_run('QR Seed Pilot Study — University Research Project')
run.font.size   = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
toc_items = [
    ('1.', 'Project Overview', '3'),
    ('2.', 'System Architecture', '4'),
    ('3.', 'Simulator 1 — BD Alaris GP Volumetric Pump', '6'),
    ('4.', 'Simulator 2 — B. Braun Infusomat Space', '10'),
    ('5.', 'Simulator 3 — Graseby 3100 Syringe Driver', '13'),
    ('6.', 'Network Simulator', '16'),
    ('7.', 'AI Risk Pipeline', '18'),
    ('8.', 'Research Panel & Task Mode', '24'),
    ('9.', 'Drug Library', '26'),
    ('10.', 'Manual References & Source Documents', '28'),
    ('11.', 'Risk Rule Reference (R01–R21)', '30'),
    ('12.', 'Device Design Score Registry', '32'),
    ('13.', 'Feature Schema — TrainingRecord', '34'),
]
for num, title, pg in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(10)
    tab_run = p.add_run(f'  .......................................  {pg}')
    tab_run.font.size  = Pt(10)
    tab_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Project Overview')

body(
    'PumpSimulator2 is a research-grade web simulator of three clinical infusion pump devices, '
    'built as the data generation engine for an AI-driven safety and security ranking system. '
    'The project is funded by a QR Seed Pilot Study: "An AI-Empowered Safety and Security '
    'Ranking System for Infusion Pump Medical Devices".'
)
doc.add_paragraph()

h2('1.1 Dual Research Goals')
body('The simulator serves two parallel objectives:')
for item in [
    ('Behavioural Fidelity',
     'Each simulated device replicates its real-world counterpart exactly as documented in the '
     'official Directions For Use (DFU) / Instructions For Use (IFU) manual. State transitions, '
     'rate limits, guardrail behaviour, and alarm conditions all source directly from manufacturer '
     'documentation.'),
    ('Data Generation',
     'Every simulator session produces a structured TrainingRecord containing 55 features across '
     'four risk layers. The simulator IS the dataset generator — real human interaction sessions '
     'are captured as labelled training data for the AI baseline model.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item[0] + ': ')
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(item[1])
    run2.font.size = Pt(10)

doc.add_paragraph()
h2('1.2 Technology Stack')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Component', 'Technology', 'Role'])
for row in [
    ('Frontend Framework', 'React 18 + TypeScript (strict)', 'UI rendering and state management'),
    ('Build Tool', 'Vite 5', 'Dev server, HMR, production build'),
    ('State Machines', 'Pure TypeScript (no React)', 'All pump business logic — pure functions'),
    ('Styling', 'Inline CSS + Share Tech Mono font', 'Retro monospace medical device aesthetic'),
    ('Data Export', 'CSV + JSON via Blob URLs', 'Training dataset download'),
    ('ID Generation', 'Web Crypto API (randomUUID)', 'Session IDs — no external packages'),
]:
    add_row(t, row, shade=row[0] in ('Frontend Framework', 'State Machines', 'Data Export'))

doc.add_paragraph()
h2('1.3 Research Basis')
for ref in [
    'Cauchi et al. (2011) "Towards Dependable Number Entry for Medical Devices" — EICS4Med Workshop, CHI-MED Project. Defines interaction features: correction_count, boundary_hit_count, golden_path_ratio.',
    'Thimbleby & Cairns (2010) "Programmable devices, interface design, and error" — J. Royal Society Interface 7(51):1429–1439. Dose error magnitude thresholds.',
    'NPSA Patient Safety Alert (2010, UK) — "Safer use of syringe drivers". Clinical basis for Graseby 3100 design risk designation.',
    'IEC 60601-2-24:2012 — International standard for infusion pump essential performance. VTBI and guardrail requirements.',
    'FDA TPLC Guidance (Dec 2014) — Infusion Pumps Total Product Life Cycle. DERS (Dose Error Reduction System) implementation guidance.',
    'CISA ICS-CERT Advisories — ICSMA-21-294-01 (B. Braun), ICSMA-23-194-01 (BD Alaris), ICSMA-22-251-01 (Baxter Sigma Spectrum). CVE data for Layer 0 design scoring.',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(ref)
    run.font.size = Pt(9)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
h1('2. System Architecture')

h2('2.1 High-Level Architecture Diagram')
code_block(
"""
┌─────────────────────────────────────────────────────────────────────────────┐
│                          PUMPSIMULATOR2  (React App)                        │
│                                                                             │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌─────────────┐  │
│  │ AlarisGP │  │  Braun   │  │ Graseby  │  │ Network  │  │  Research   │  │
│  │   Tab    │  │Infusomat │  │  3100    │  │Simulator │  │   Panel     │  │
│  │          │  │   Tab    │  │   Tab    │  │   Tab    │  │    Tab      │  │
│  └────┬─────┘  └────┬─────┘  └────┬─────┘  └────┬─────┘  └──────┬──────┘  │
│       │              │              │              │               │         │
│  ┌────▼─────┐  ┌────▼─────┐  ┌────▼─────┐        │         ┌────▼──────┐  │
│  │ Pump     │  │ Braun    │  │ Graseby  │  ┌──────▼──────┐  │ Task Mode │  │
│  │ Context  │  │ Pump     │  │ Pump     │  │  Network    │  │ Scenario  │  │
│  │ (React)  │  │ Context  │  │ Context  │  │  Context    │  │  Runner   │  │
│  └────┬─────┘  └────┬─────┘  └────┬─────┘  └──────┬──────┘  └────┬──────┘  │
│       │              │              │               │               │        │
│  ┌────▼─────┐  ┌────▼─────┐  ┌────▼─────┐  ┌──────▼──────┐  ┌────▼──────┐  │
│  │ usePump  │  │useBraun  │  │useGraseby│  │ useNetwork  │  │sessionAd  │  │
│  │  Hook    │  │  Pump    │  │  Pump    │  │  (hook)     │  │  apter.ts │  │
│  └────┬─────┘  └────┬─────┘  └────┬─────┘  └──────┬──────┘  └────┬──────┘  │
│       │              │              │               │               │        │
│  ┌────▼─────┐  ┌────▼─────┐  ┌────▼─────┐  ┌──────▼──────┐  ┌────▼──────┐  │
│  │state     │  │ braun    │  │ graseby  │  │ connection  │  │ extract   │  │
│  │Machine   │  │ State    │  │ State    │  │ Machine.ts  │  │ Features  │  │
│  │  .ts     │  │ Machine  │  │ Machine  │  │ packet      │  │ labelling │  │
│  │(pure TS) │  │  .ts     │  │  .ts     │  │ Generator   │  │ Rules     │  │
│  └──────────┘  └──────────┘  └──────────┘  └─────────────┘  └───────────┘  │
│                                                                             │
│  ══════════════════════════ AI PIPELINE ════════════════════════════════    │
│  featureExtractor.ts → labellingRules.ts → deviceDesign.ts → TrainingRecord │
└─────────────────────────────────────────────────────────────────────────────┘
"""
)
doc.add_paragraph()

h2('2.2 Layer Separation Principle')
body(
    'The architecture enforces a strict separation between pure logic and React UI. '
    'All pump state machines, AI pipeline code, and network logic live in src/pump/, '
    'src/ai/, and src/network/ respectively. These files contain zero React imports. '
    'React hooks wrap the pure functions, and React components only call hooks and render.'
)
doc.add_paragraph()

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
table_header(t, ['Directory', 'React?', 'Contents', 'Rule'])
for row in [
    ('src/pump/', 'NO', 'State machines, types, constants, drug library, guardrails, alarms', 'Pure TypeScript only'),
    ('src/pump/braun/', 'NO', 'B. Braun state machine, types, constants, drug library, guardrails', 'Pure TypeScript only'),
    ('src/pump/graseby/', 'NO', 'Graseby 3100 state machine, types, constants', 'Pure TypeScript only'),
    ('src/ai/', 'NO', 'Feature extractor, labelling rules, device design, scenario generator, session adapter', 'Pure TypeScript only'),
    ('src/network/', 'NO', 'Network types, connection machine, packet generator', 'Pure TypeScript only'),
    ('src/hooks/', 'YES', 'usePump, useBraunPump, useGrasebyPump, useLogger, useHoldRepeat', 'React hooks only'),
    ('src/contexts/', 'YES', 'PumpContext, BraunPumpContext, GrasebyPumpContext, NetworkContext', 'React contexts'),
    ('src/components/', 'YES', 'AlarisGP, BraunInfusomat, Graseby3100, NetworkSimulator, ResearchPanel', 'React components'),
]:
    add_row(t, row, shade=row[1] == 'NO')

doc.add_paragraph()
h2('2.3 Data Flow — Session to TrainingRecord')
code_block(
"""
User Interaction (chevron presses, button clicks)
        │
        ▼
  React Component (AlarisGP.tsx / BraunInfusomat.tsx / Graseby3100.tsx)
        │  calls handlers
        ▼
  React Hook (usePump / useBraunPump / useGrasebyPump)
        │  dispatches to pure state machine
        ▼
  Pure State Machine (stateMachine.ts / braunStateMachine.ts / grasebyStateMachine.ts)
        │  returns { state: PumpState, logEntries: SessionLogEntry[] }
        ▼
  Session Logger (useLogger / useBraunLogger / useGrasebyLogger)
        │  accumulates immutable log entries
        ▼
  [infusion_started event triggers capture in Task Mode]
        │
        ▼
  sessionAdapter.ts → buildTrainingRecord()
        │  normalises device-specific types to Alaris GP common types
        │  calls extractFeatures() → applyLabellingRules() → computeDesignScore()
        ▼
  TrainingRecord (55 fields, 4-layer scores, A+–F grade)
        │
        ▼
  ResearchPanel (display) + CSV / JSON export
"""
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. BD ALARIS GP
# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Simulator 1 — BD Alaris GP Volumetric Pump')

h2('3.1 Device Overview')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
table_header(t, ['Property', 'Value'])
for row in [
    ('Device Type', 'Large Volume Volumetric Infusion Pump'),
    ('Manufacturer', 'BD (Becton, Dickinson & CareFusion)'),
    ('Model Identifier', 'alaris_gp'),
    ('Firmware (simulator)', '9.12'),
    ('Rate Range', '0.1 – 1200.0 ml/h'),
    ('VTBI Range', '0.1 – 9999 ml'),
    ('KVO Rate', '1.0 ml/h (after VTBI complete)'),
    ('Bolus Rate Default', '500 ml/h (max 1200 ml/h)'),
    ('Bolus Volume Max', '5 ml'),
    ('Pressure Levels', 'L0 – L8 (default L5)'),
    ('Drug Library', 'Yes — 10 drugs + MANUAL'),
    ('Guardrails', 'Soft limits (warning) + Hard limits (blocked)'),
    ('VTBI', 'Optional — required for auto-stop and KVO'),
    ('Network', 'Yes — BD Alaris Gateway (simulated)'),
    ('Anti-freeflow', 'Yes'),
    ('KVO', 'Yes — 1.0 ml/h after VTBI completion'),
]:
    add_row(t, row, shade=row[0] in ('Rate Range', 'Drug Library', 'Guardrails', 'Network'))

doc.add_paragraph()
h2('3.2 Manual References')
for ref in [
    'BD document 1000DF00152 Issue 1 — Primary DFU (Directions For Use)',
    'BD document BDDF00535 Issue 4 — DFU with Guardrails specification',
    'BD 1000SM00013 Issue 4 — Technical Service Manual',
    'PVSio-web formal model: github.com/pvsioweb/pvsio-web (AlarisGP demo)',
    'CISA ICSMA-23-194-01 (2023) — BD Alaris System CVEs (13 vulnerabilities, max CVSS 9.8)',
    'FDA MedWatch — 3 Class I recalls, 2 Class II recalls documented in device design profile',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(ref)
    run.font.size = Pt(9)

doc.add_paragraph()
h2('3.3 Screen State Machine')
code_block(
"""
  LANGUAGE_SELECT
        │ language chosen
        ▼
  DRUG_SELECT  ◄──────────────────────────────────────────────────────────────┐
        │ drug chosen                                                          │
        ▼                                                                      │
  WEIGHT_ENTRY (weight-based drugs only)                                      │
        │ weight confirmed / skipped                                           │
        ▼                                                                      │
  RATE_ENTRY  ◄────────────── ON_HOLD (RE-PROG softkey)                       │
        │ RUN pressed:                                                         │
        ├──[rate in hard limits]──► GUARDRAIL_BLOCKED ──► RATE_ENTRY          │
        ├──[rate in soft limits warning]──► GUARDRAIL_WARNING                 │
        │       ├──[OVERRIDE]──► RUNNING                                      │
        │       └──[RE-ENTER]──► RATE_ENTRY                                   │
        └──[rate OK]──────────────► RUNNING ◄──── ON_HOLD (RUN)               │
                                      │ HOLD                                  │
                                      ▼                                       │
                                   ON_HOLD ──────────────────────────────────►│
                                      │ alarm                                 │
                                      ▼                                       │
                                    ALARM ──[SILENCE]──► ON_HOLD              │
                                                                              │
  OPTIONS ◄────────── RUNNING / ON_HOLD (OPTIONS button)                      │
       └──► VTBI_ENTRY / PRESSURE_VIEW                                        │
                                                                              │
  Any screen ──[Power hold 3s]──► LANGUAGE_SELECT (full reset) ──────────────┘
"""
)

doc.add_paragraph()
h2('3.4 Controls — DFU Manual Reference')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
table_header(t, ['Control', 'DFU Code', 'Manual Description', 'Simulator Behaviour'])
for row in [
    ('ON/OFF', 'a', 'Press once ON; hold 3s OFF', 'Click = ON; hold 3s → confirm reset dialog'),
    ('RUN', 'b', 'Start infusion; green LED flashes', 'Starts infusion; triggers guardrail check'),
    ('HOLD', 'h', 'Put infusion on hold; amber LED lit', 'Pauses infusion; amber indicator'),
    ('MUTE', 'c', 'Silence alarm ~2 min; resounds after', 'Silences for 120s; logs mute_pressed event'),
    ('BOLUS', 'i', 'Hold softkey to operate; adds to volume', 'Hold-to-deliver; stops on release'),
    ('OPTIONS', 'd', 'Access optional features', 'Opens options menu'),
    ('PRESSURE', 'e', 'Display pressure and adjust alarm limit', 'Shows pressure bar L0–L8; adjustable'),
    ('«« « » »»', 'f', 'Double=faster, Single=slower adjust', '«« = -10, « = -1, » = +1, »» = +10'),
    ('SOFTKEYS', 'g', 'Context-sensitive; 3 total', 'Change per screen (OVERRIDE / RE-ENTER etc.)'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('3.5 Chevron Hold-to-Accelerate Logic')
code_block(
"""
  mousedown / touchstart:
    → immediate single step applied (first press)
    → 500ms timer started

  After 500ms hold:
    → fires every 80ms until mouseup / touchend / mouseleave

  Rate clamping (ABSOLUTE, no stateful memory):
    rate + delta > RATE_MAX  →  rate = RATE_MAX  +  log boundary_hit
    rate + delta < RATE_MIN  →  rate = RATE_MIN  +  log boundary_hit

  Source: DFU Section 3 "faster/slower" description
  Constants: STEP_LARGE=10, STEP_SMALL=1, HOLD_DELAY_MS=500, HOLD_REPEAT_MS=80
"""
)

doc.add_paragraph()
h2('3.6 Guardrail System')
body(
    'The Alaris GP implements a two-tier Dose Error Reduction System (DERS). '
    'Soft limits produce a warning that the operator can override (logged as guardrail_override). '
    'Hard limits block the infusion entirely and force re-entry (logged as guardrail_blocked). '
    'Source: DFU BDDF00535 Issue 4 — Guardrails specification.'
)
code_block(
"""
  checkGuardrail(rate, drug):
    if drug.id === "manual" → { status: "ok" }           // no guardrails in MANUAL mode
    if rate > drug.hardMax || rate < drug.hardMin → { status: "blocked" }
    if rate > drug.softMax || rate < drug.softMin → { status: "warning", message: "..." }
    else → { status: "ok" }

  Screen transitions triggered:
    "blocked"  → GUARDRAIL_BLOCKED  (only option: RE-ENTER → RATE_ENTRY)
    "warning"  → GUARDRAIL_WARNING  (OVERRIDE → RUNNING | RE-ENTER → RATE_ENTRY)
    "ok"       → RUNNING
"""
)

doc.add_paragraph()
h2('3.7 Alarm System')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Alarm Type', 'Trigger Condition', 'Effect'])
for row in [
    ('OCCLUSION', 'pressureLevel >= 7 during running', 'Stops infusion → ALARM screen'),
    ('AIR_IN_LINE', 'volumeInfused >= 500ml (auto) or research button', 'Stops infusion → ALARM screen'),
    ('INFUSION_COMPLETE', 'volumeInfused >= vtbi', 'KVO mode → rate drops to 1.0 ml/h'),
    ('BATTERY_LOW', 'batteryLevel < 15%', 'Warning alarm; continues running'),
    ('AC_FAIL', 'network_connected = false AND battery mode', 'Warning alarm'),
    ('RATE_TOO_HIGH', 'Guardrail soft max exceeded (advisory)', 'Advisory only; does not stop infusion'),
    ('RATE_TOO_LOW', 'Guardrail soft min exceeded (advisory)', 'Advisory only'),
    ('KVO', 'VTBI complete → rate auto-drops', 'Informational; running at KVO rate'),
]:
    add_row(t, row)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. B. BRAUN INFUSOMAT
# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Simulator 2 — B. Braun Infusomat Space')

h2('4.1 Device Overview')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
table_header(t, ['Property', 'Value'])
for row in [
    ('Device Type', 'Large Volume Volumetric Infusion Pump'),
    ('Manufacturer', 'B. Braun Melsungen AG'),
    ('Model Identifier', 'braun_infusomat'),
    ('Firmware (simulator)', '3.10'),
    ('Rate Range', '0.1 – 999.9 ml/h'),
    ('VTBI Range', '0.1 – 9999 ml'),
    ('KVO Rate', '1.0 ml/h'),
    ('Bolus Volume Max', '2 ml (smaller than Alaris GP)'),
    ('Drug Library', 'Yes — dedicated B. Braun drug library (10 drugs)'),
    ('Guardrails', '3-tier: Advisory + Soft warning + Hard blocked'),
    ('SpaceCom2 Network', 'Yes — B. Braun SpaceStation protocol (simulated)'),
    ('AIL Sensitivity', 'Configurable 20–500 µl'),
    ('Patient Weight', 'Optional (1–200 kg) for weight-based dosing'),
    ('KVO', 'Yes — 1.0 ml/h after VTBI completion'),
]:
    add_row(t, row, shade=row[0] in ('Rate Range', 'Drug Library', 'Guardrails', 'SpaceCom2 Network'))

doc.add_paragraph()
h2('4.2 Manual References')
for ref in [
    'B. Braun Infusomat Space IFU (686N GB) — braun.com (included in docs/braun_infusomat_space_ifu_686N_GB.pdf)',
    'B. Braun Infusomat Space IFU (586U US) — bbraunusa.com (included in docs/braun_infusomat_space_ifu_586U_US.pdf)',
    'B. Braun Infusomat Space Service Manual — included in docs/braun_infusomat_service_manual.pdf',
    'CISA ICSMA-21-294-01 (2021) — 5 CVEs including CVE-2021-33885 (CVSS 9.0): unauthenticated firmware modification via SpaceStation',
    'CVE-2021-33883 (CVSS 7.1) — cleartext transmission of sensitive data over SpaceStation protocol',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(ref)
    run.font.size = Pt(9)

doc.add_paragraph()
h2('4.3 Key Differentiators from Alaris GP')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Feature', 'Alaris GP', 'B. Braun Infusomat Space'])
for row in [
    ('Guardrail tiers', '2 (soft + hard)', '3 (advisory + soft + hard) — safer design'),
    ('Display type', 'Text LCD', 'Graphic LCD — better visual confirmation'),
    ('Bolus max', '5 ml', '2 ml — lower dose error amplification'),
    ('Network protocol', 'BD Alaris Network', 'SpaceStation SpaceCom2 protocol'),
    ('CVE count', '13 (CVSS 9.8)', '5 (CVSS 9.0)'),
    ('Drug cursor nav', 'Softkey-based', 'Dedicated UP/DOWN buttons'),
    ('VTBI entry', 'Via OPTIONS menu', 'Direct VTBI softkey on main screen'),
    ('Patient weight', 'Set once at session start', 'Stored per drug profile'),
]:
    add_row(t, row, shade=row[0] in ('Guardrail tiers', 'CVE count'))

doc.add_paragraph()
h2('4.4 Screen State Machine')
code_block(
"""
  BOOT  →  DRUG_SELECT
               │ drug + weight (if needed) confirmed
               ▼
          VTBI_ENTRY (optional)
               │
               ▼
          RATE_ENTRY  ◄──────── ON_HOLD (RE-PROG)
               │ RUN:
               ├──[hard limit]──► GUARDRAIL_BLOCKED ──► RATE_ENTRY
               ├──[soft limit]──► GUARDRAIL_WARNING
               │     ├──[OVERRIDE]──► RUNNING
               │     └──[RE-ENTER]──► RATE_ENTRY
               └──[advisory]───► ADVISORY_NOTICE ──► RUNNING (auto-dismissed)
                                        │ (unique to B. Braun — advisory tier)
               [rate OK]──────────────► RUNNING
                                          │ HOLD
                                          ▼
                                       ON_HOLD ──[START]──► RUNNING
                                          │ alarm
                                          ▼
                                        ALARM ──[SILENCE]──► ON_HOLD
"""
)

doc.add_paragraph()
h2('4.5 SpaceCom2 Network Simulation')
body(
    'The B. Braun network connectivity is simulated through the spacecom2Connected flag in '
    'BraunPumpState. When the NetworkSimulator tab activates the SpaceStation connection, '
    'this flag syncs to the NetworkContext, which is then read by the Task Mode session '
    'adapter to auto-populate the network_connected field in the TrainingRecord.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. GRASEBY 3100
# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Simulator 3 — Graseby 3100 Syringe Driver')

h2('5.1 Device Overview')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
table_header(t, ['Property', 'Value'])
for row in [
    ('Device Type', 'Ambulatory Syringe Driver'),
    ('Manufacturer', 'Graseby Medical Ltd (now Smiths Medical)'),
    ('Model Identifier', 'graseby_3100'),
    ('Firmware (simulator)', '3100-v1.0'),
    ('Rate Range', '0.1 – 199.9 ml/h in 0.1 ml/h steps'),
    ('Supported Syringe Sizes', '20 ml / 30 ml / 50 ml'),
    ('Volume Counter', '0.0 – 999.9 ml (running total)'),
    ('VTBI', 'NONE — syringe capacity is the only volume limit'),
    ('Drug Library', 'NONE — always MANUAL mode'),
    ('Guardrails', 'NONE — any rate between 0.1 and 199.9 is accepted'),
    ('Network', 'NONE — fully standalone device'),
    ('Bolus Mode', 'NONE'),
    ('KVO', 'NONE — infusion ends on alarm when syringe empty'),
    ('Anti-freeflow', 'NONE'),
    ('Battery', 'Sealed lead acid rechargeable (simulated as 0–100%)'),
    ('Drive Accuracy', '±2% (per Technical Service Manual 00SM-0131-7)'),
]:
    add_row(t, row, shade=row[0] in ('VTBI', 'Drug Library', 'Guardrails', 'Network'))

doc.add_paragraph()
h2('5.2 Manual References')
for ref in [
    'Graseby 3100 Technical Service Manual Part No. 00SM-0131-7, Issue 7, August 2004 — Smiths Medical International Ltd. (included in docs/graseby_3100_technical_service_manual_2004.pdf)',
    'Graseby 3100 Operators Manual — ardusmedical.com (2002, scanned image)',
    'NPSA Patient Safety Alert — "Safer use of syringe drivers in palliative care" (2010, UK). Documents fatal overdoses linked to absence of guardrails.',
    'Graseby 3300 PCA Technical Service Manual Part No. 00SM-0128-7 — included in docs/graseby_3300_pca_technical_service_manual_2004.pdf',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(ref)
    run.font.size = Pt(9)

doc.add_paragraph()
h2('5.3 Key Specification — Source: 00SM-0131-7 (confirmed)')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Spec', 'Manual Value', 'Simulator Implementation'])
for row in [
    ('Flow rate range', '0.1 to 199.9 ml/hour in 0.1 ml increments', 'RATE_MIN=0.1, RATE_MAX=199.9 (grasebyConstants.ts)'),
    ('Syringe sizes', 'BD Plastipak / Braun Omnifix / Monoject: 20, 30, 50/60 ml', 'SyringeCapacityMl = 20 | 30 | 50 (grasebyTypes.ts)'),
    ('Volume counter', '0 to 999.9 ml in 0.1 ml increments', 'volumeInfused field (continuous accumulation)'),
    ('Occlusion pressure', 'Adjustable 250–600 mmHg', 'Not modelled — research simplification'),
    ('Drive accuracy', '±2%', 'Not modelled — deterministic simulation'),
    ('Battery', 'Sealed lead acid, rechargeable (>3 hours)', 'Abstract 0–100%, drains per tick'),
    ('Syringe sensing', 'Automatic (size sensor flag)', 'Manual picker — UI research simplification'),
    ('Totaliser', 'Volume accumulation display', 'volumeInfused shows running total'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('5.4 Screen State Machine')
code_block(
"""
  BOOT  →  RATE_ENTRY  (device powers on, self-test complete)
               │
               │  [Chevron presses adjust rate — no guardrail check]
               │
               │  START pressed (rate > 0):
               ▼
            RUNNING  ──────────────────────────► ALARM (syringe_empty / occlusion / battery_low)
               │ STOP                                │ ALARM SILENCE
               ▼                                     │ OCCLUSION/BATTERY_LOW → ON_HOLD
            ON_HOLD                                  │ SYRINGE_EMPTY         → RATE_ENTRY
               │ START → RUNNING                     ▼
               │ RE-PROGRAM → RATE_ENTRY
               ▼
           RATE_ENTRY  (re-enter rate for new syringe)

  Any screen ──[Power Off]──► BOOT (full reset to initial state)
"""
)

doc.add_paragraph()
h2('5.5 Clinical Risk Context (Research Rationale)')
body(
    'The Graseby 3100 is the "baseline high design risk" device in the ranking system. '
    'Its absence of drug library, guardrails, and VTBI means every session triggers '
    'at minimum R12 (no library, MEDIUM) and R13 (no VTBI, MEDIUM), placing every '
    'session at MEDIUM risk regardless of operator behaviour. This creates a useful '
    'three-way comparison:'
)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Device', 'Design Risk Profile', 'Cyber Risk Profile'])
for row in [
    ('BD Alaris GP', 'MEDIUM — has guardrails, VTBI, drug library', 'HIGH — 13 CVEs, CVSS 9.8, unsigned firmware'),
    ('B. Braun Infusomat Space', 'LOW — 3-tier guardrails, graphic LCD, anti-freeflow', 'HIGH — 5 CVEs, CVSS 9.0, unsigned firmware'),
    ('Graseby 3100', 'HIGH — no guardrails, no library, no VTBI', 'ZERO — pre-network era, no software'),
]:
    add_row(t, row)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. NETWORK SIMULATOR
# ═══════════════════════════════════════════════════════════════════════════════
h1('6. Network Simulator')

h2('6.1 Overview')
body(
    'The Network Simulator tab models the BD Alaris pump communicating with a hospital '
    'Gateway and Drug Library Server over a simulated 802.11 wireless network. '
    'It demonstrates real-world attack vectors documented in published CVEs for the BD Alaris system. '
    'The simulation state auto-syncs to the NetworkContext, which feeds the network_connected '
    'field (Layer 3) into all Alaris GP TrainingRecords.'
)

doc.add_paragraph()
h2('6.2 Network State Machine')
code_block(
"""
  OFFLINE
     │ Connect initiated
     ▼
  SCANNING  (searching for AP)
     │
     ▼
  ASSOCIATING  (joining SSID)
     │
     ▼
  DHCP  (acquiring IP address)
     │
     ▼
  AUTHENTICATING  (pump ↔ Gateway auth exchange)
     │
     ▼
  CONNECTED  (heartbeat packets flowing: PUMP→GW every ~2s)
     │
     ▼
  SYNCING_LIBRARY  (LIBRARY_REQUEST → LIBRARY_RESPONSE)
     │
     ▼
  LIBRARY_CURRENT  (drug library up to date)
     │
     └── [attack scenario triggered] → Attack packets injected into packet log
"""
)

doc.add_paragraph()
h2('6.3 Network Nodes')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Node', 'Type', 'Role'])
for row in [
    ('Alaris GP Pump', 'pump', 'Sends heartbeat, alarm, infusion data; requests/receives drug library'),
    ('Access Point (AP)', 'ap', 'Wireless relay — PUMP↔GW bridge'),
    ('Alaris Gateway', 'gateway', 'Central clinical hub — logs alarms, distributes drug libraries'),
    ('Drug Library Server', 'server', 'Hosts authoritative drug library; responds to LIBRARY_REQUEST'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('6.4 Packet Types')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Packet Type', 'Direction', 'Description'])
for row in [
    ('HEARTBEAT', 'PUMP→GW', 'Periodic liveness signal (every ~2s)'),
    ('ALARM_EVENT', 'PUMP→GW', 'Alarm triggered — clinical notification'),
    ('INFUSION_DATA', 'PUMP→GW', 'Rate, volume, drug, patient weight'),
    ('LIBRARY_REQUEST', 'PUMP→GW', 'Request current drug library version'),
    ('LIBRARY_RESPONSE', 'GW→PUMP / SERVER→GW', 'Drug library payload delivery'),
    ('FIRMWARE_CHECK', 'PUMP→GW', 'Firmware version verification request'),
    ('FIRMWARE_RESPONSE', 'GW→PUMP', 'Firmware update / version confirmed'),
    ('ACK / NACK', 'GW→PUMP', 'Message acknowledgement / rejection'),
    ('MITM_INJECT', 'ATTACKER→GW', 'ATTACK: Modified drug library injected in transit'),
    ('REPLAY_ATTACK', 'ATTACKER→GW', 'ATTACK: Old library version replayed to pump'),
    ('SPOOFED_ACK', 'ATTACKER→PUMP', 'ATTACK: Forged ACK bypasses authentication'),
    ('FIRMWARE_INJECT', 'ATTACKER→PUMP', 'ATTACK: Unsigned CVE-affected firmware pushed'),
]:
    add_row(t, row, shade='ATTACK' in row[2])

doc.add_paragraph()
h2('6.5 Attack Scenarios (CVE-Mapped)')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
table_header(t, ['Scenario', 'CVE Reference', 'Severity', 'Effect on Pump'])
for row in [
    ('MITM Drug Library', 'CVE-2020-25165', 'CRITICAL', 'Guardrail limits silently changed — pump accepts dangerous rates'),
    ('Replay Old Library', '—', 'HIGH', 'Drug library reverted to older version with wider limits'),
    ('Spoofed ACK', 'CVE-2020-25163', 'HIGH', 'Alarm silently lost — clinical staff not notified'),
    ('Firmware Injection', '—', 'CRITICAL', 'Pump firmware replaced with CVE-affected version 8.05'),
]:
    add_row(t, row, shade=row[2] == 'CRITICAL')

doc.add_paragraph()
h2('6.6 IDS — Intrusion Detection Rules')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Rule ID', 'Rule Name', 'Detection Logic'])
for row in [
    ('IDS-R01', 'MITM Drug Library Substitution', 'Packet type === MITM_INJECT'),
    ('IDS-R02', 'Replay Attack — Old Library Version', 'Packet type === REPLAY_ATTACK'),
    ('IDS-R03', 'Spoofed ACK — Authentication Bypass', 'Packet type === SPOOFED_ACK'),
    ('IDS-R04', 'Firmware Injection Attempt', 'Packet type === FIRMWARE_INJECT'),
    ('IDS-R05', 'Malicious Packet Source', 'Direction includes ATTACKER→'),
    ('IDS-R06', 'Heartbeat Flood / DoS Pattern', '>5 HEARTBEAT packets in 3 seconds'),
    ('IDS-R07', 'Unexpected Library Change During Infusion', 'LIBRARY_RESPONSE within 5s of INFUSION_DATA'),
]:
    add_row(t, row)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. AI RISK PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════
h1('7. AI Risk Pipeline')

h2('7.1 Four-Layer Composite Risk Model')
body(
    'The core output of the simulator is a 4-layer composite risk score (0–1) and an '
    'energy-label style grade (A+ through F). Each layer captures a distinct clinical or '
    'technical risk dimension. The composite formula weights the layers as follows:'
)
code_block(
"""
  composite_score = 0.20 × design_score        (Layer 0 — deviceDesign.ts)
                  + 0.30 × interaction_score    (Layer 1 — labellingRules.ts)
                  + 0.25 × configuration_score  (Layer 2 — labellingRules.ts)
                  + 0.25 × system_score         (Layer 3 — labellingRules.ts)

  Grade thresholds:
    0.00 – 0.10  =  A+   →  risk_label = "low"
    0.11 – 0.20  =  A    →  risk_label = "low"
    0.21 – 0.35  =  B    →  risk_label = "low"
    0.36 – 0.50  =  C    →  risk_label = "medium"
    0.51 – 0.65  =  D    →  risk_label = "medium"
    0.66 – 0.80  =  E    →  risk_label = "high"
    0.81 – 1.00  =  F    →  risk_label = "high"
"""
)

doc.add_paragraph()
h2('7.2 Layer 0 — Device Design Score (deviceDesign.ts)')
body(
    'Fixed per device model. Captures how safely the device was designed, independent '
    'of how it is used in any individual session. Computed once at module load time from '
    'the DeviceDesignProfile registry using penalty weights.'
)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Design Factor', 'Penalty Weight', 'Source'])
for row in [
    ('No drug library (DERS absent)', '+0.30', 'FDA TPLC Guidance Dec 2014'),
    ('Guardrail tiers=0 (no guardrails)', '+0.25', 'IEC 60601-2-24:2012'),
    ('Guardrail tiers=1 (hard only)', '+0.10', 'IEC 60601-2-24:2012'),
    ('Guardrail tiers=2 (soft+hard)', '+0.05', 'IEC 60601-2-24:2012'),
    ('No VTBI (indefinite infusion risk)', '+0.15', 'IEC 60601-2-24:2012'),
    ('No anti-freeflow protection', '+0.10', 'Generic Infusion Pump Hazard Analysis H-12'),
    ('Per CVE (×cve_count, cap 0.40)', '+0.04 each', 'CISA ICS-CERT advisories'),
    ('Critical CVSS score (>=9.0)', '+0.15', 'NVD CVSS v3.1 severity bands'),
    ('Firmware not cryptographically signed', '+0.15', 'CVE-2021-33885 precedent'),
    ('Cleartext network transmission', '+0.08', 'CVE-2021-33883, CVE-2022-26390'),
    ('Per FDA Class I recall (cap 0.24)', '+0.06 each', 'FDA MedWatch'),
    ('Per FDA Class II recall (cap 0.08)', '+0.02 each', 'FDA MedWatch'),
    ('Bolus max > 5 ml', '+0.08', 'Clinical dose error amplification'),
    ('Display quality: basic_numeric', '+0.08', 'FDA Human Factors Guidance 2011'),
    ('Display quality: text_lcd', '+0.04', 'FDA Human Factors Guidance 2011'),
    ('Display quality: graphic_lcd', '+0.01', 'FDA Human Factors Guidance 2011'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('7.3 Design Score Summary — All Registered Devices')
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
table_header(t, ['Device', 'CVEs', 'Max CVSS', 'Key Gaps', 'Approx. Design Score'])
for row in [
    ('B. Braun Infusomat Space', '5', '9.0', 'Unsigned FW, cleartext TX, 2 Class I recalls', '~0.57 (D)'),
    ('BD Alaris GP', '13', '9.8', 'Unsigned FW, cleartext TX, 3 Class I recalls', '~0.75 (E)'),
    ('Graseby 3100', '0', '0', 'No library, no guardrails, no VTBI, no anti-freeflow', '~0.82 (F)'),
    ('ICU Medical Plum 360', '0', '0', 'Unsigned FW, 2 Class I recalls', '~0.27 (B)'),
    ('Baxter Sigma Spectrum', '4', '7.5', 'Unsigned FW, cleartext TX, 3 Class I recalls', '~0.59 (D)'),
    ('Fresenius Agilia VP', '13', '7.5', 'Unsigned FW, cleartext TX, 1 Class I recall', '~0.73 (E)'),
    ('CADD-Solis PCA', '2', '9.9', 'No anti-freeflow, no KVO, Critical CVE, 3 Class I recalls', '~0.82 (F)'),
]:
    add_row(t, row, shade=row[0] in ('Graseby 3100', 'BD Alaris GP'))

doc.add_paragraph()
h2('7.4 Layer 1 — Interaction Score (nurse programming behaviour)')
body(
    'Derived from session log events. Captures how the operator programmed the rate: '
    'error magnitude, corrections, boundary hits, entry speed. '
    'Source: Cauchi et al. (2011) EICS4Med — interaction error taxonomy.'
)

h2('7.5 Layer 2 — Configuration Score (session setup decisions)')
body(
    'Derived from the pump setup state at the moment infusion started. Captures decisions '
    'like whether VTBI was set, whether drug library was used, whether guardrails were overridden. '
    'Source: ISMP High-Alert Medications; IEC 60601-2-24:2012.'
)

h2('7.6 Layer 3 — System Score (device physical/cyber state)')
body(
    'Derived from UIDeviceContext (researcher-entered) or ScenarioProfile (synthetic). '
    'Captures maintenance history, battery, network connectivity, drug library age, '
    'firmware version, and recent alarm history. '
    'Source: FDA Infusion Pump Safety guidance; CISA ICS-CERT advisories.'
)

doc.add_paragraph()
h2('7.7 Scenario Profiles (Synthetic Dataset Generation)')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
table_header(t, ['Profile', 'Weight', 'Device State', 'Expected Risk Distribution'])
for row in [
    ('ideal', '5', 'FW=9.12, maint 0-30d, batt 80-100%, connected, library 0-14d', 'Mostly LOW/MEDIUM'),
    ('neglected', '2', 'FW=7.2.0, maint 90-400d, batt 15-50%, disconnected, library 91-365d', 'Mostly MEDIUM'),
    ('cyber_risk', '1', 'FW=random CVE version, connected, well maintained', 'All HIGH (R05 fires)'),
    ('emergency', '2', 'FW=9.12, MANUAL mode only, high error probability, no VTBI', 'HIGH/MEDIUM mix'),
]:
    add_row(t, row)
note('Weights are sampling weights for dataset generation. Total sessions = sum(weight × n_per_profile).')

doc.add_paragraph()
h2('7.8 Feature Extraction — Key Metrics')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Feature', 'Formula / Source', 'Purpose'])
for row in [
    ('golden_path_ratio', 'actual_keypresses / minimum_possible_keypresses', 'Efficiency of rate entry — 1.0 = perfect. Cauchi et al. (2011)'),
    ('relative_error', 'abs(final - intended) / intended (capped 2.0)', 'Dose error magnitude. Thimbleby & Cairns (2010)'),
    ('large_btn_ratio', '(large_up + large_down) / total_keypresses', 'Preference for coarse vs fine adjustments'),
    ('confirmed_incorrect', '1 if relative_error > 25% AND infusion started', 'Hard indicator of started infusion with wrong dose'),
    ('entry_time_ms', 'first rate_adjust → infusion_started', 'Programming speed — >60s triggers R21'),
    ('correction_count', 'Direction reversal events in session log', 'Hesitation / uncertainty during entry'),
    ('boundary_hit_count', 'Rate went to RATE_MIN or RATE_MAX', 'Operator tested limits — triggers R15'),
]:
    add_row(t, row)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. RESEARCH PANEL
# ═══════════════════════════════════════════════════════════════════════════════
h1('8. Research Panel & Task Mode')

h2('8.1 Research Panel Tabs')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
table_header(t, ['Tab', 'Function'])
for row in [
    ('Task Mode', 'Researcher sets target rate, operator uses any of the 3 simulators. infusion_started event triggers buildTrainingRecord() and captures the session as a labelled TrainingRecord.'),
    ('Scenario Runner', 'Generates N synthetic TrainingRecords using SCENARIO_PROFILES. Runs algorithmically without UI interaction. Configurable count (10–1000) and device filter (Alaris, Braun, Graseby, Combined).'),
    ('Unified Dataset', 'Merged view of task-mode records + scenario records with device/risk/source filters. Shows grade distribution chart, export to CSV/JSON.'),
    ('Metrics', 'Live session metrics for the current simulator: keypresses, corrections, boundary hits, override count, golden_path_ratio.'),
    ('AI Rules', 'Displays all R01–R21 rules with their current evaluation state against the live session.'),
    ('Dataset Inspector', 'Tabular view of all saved TrainingRecords with column sorting and risk filter.'),
]:
    add_row(t, row, shade=row[0] in ('Task Mode', 'Unified Dataset'))

doc.add_paragraph()
h2('8.2 Task Mode Device Selection')
body(
    'Task Mode supports all three simulators simultaneously. The researcher selects which device '
    'the operator should use (Alaris GP, B. Braun, or Graseby). The corresponding pump context '
    'is watched for infusion_started events. When fired, the session log and final pump state '
    'are passed to buildTrainingRecord() via sessionAdapter.ts.'
)

doc.add_paragraph()
h2('8.3 UIDeviceContext — Researcher-Configurable Layer 3 Fields')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Field', 'Default', 'Effect on Risk Score'])
for row in [
    ('days_since_maintenance', '30', '>365 days → R06 HIGH'),
    ('battery_level_pct', '100', '<20% → R17 MEDIUM'),
    ('network_connected', '1 (Alaris: from NetworkContext; Braun: from SpaceCom2)', '0 → R18 MEDIUM'),
    ('drug_library_age_days', '14', '>90 days → R16 MEDIUM'),
    ('config_drift_score', '0.05', '>0.5 → R19 MEDIUM'),
    ('recent_occlusion_alarms', '0', '≥3 → R20 MEDIUM'),
]:
    add_row(t, row)

note('network_connected is auto-populated from live NetworkContext (Alaris) or spacecom2Connected (B. Braun) or forced 0 (Graseby).')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. DRUG LIBRARY
# ═══════════════════════════════════════════════════════════════════════════════
h1('9. Drug Library')

h2('9.1 Alaris GP Drug Library (drugLibrary.ts)')
body(
    'The Alaris GP drug library contains 10 entries (including MANUAL mode). '
    'All values are clinically validated. Hard and soft limits are sourced from '
    'clinical pharmacology references. The library cannot be changed — any modification '
    'would compromise research data integrity.'
)
t = doc.add_table(rows=1, cols=6)
t.style = 'Table Grid'
table_header(t, ['Drug', 'Unit', 'Conc.', 'Soft Min/Max', 'Hard Min/Max', 'Clinical Context'])
for row in [
    ('MANUAL ml/h', 'ml/h', '—', '1 / 1200', '0.1 / 1200', 'Direct rate entry — highest risk (no guardrails)'),
    ('ADRENALINE', 'µg/kg/min', '0.08 mg/ml', '0.01 / 0.5', '0.001 / 1.0', 'Vasopressor — narrow therapeutic window, ICU'),
    ('MORPHINE', 'mg/h', '1 mg/ml', '1 / 10', '0.5 / 20', 'Opioid — respiratory depression risk above soft max'),
    ('HEPARIN', 'U/h', '1000 U/ml', '500 / 2000', '100 / 5000', 'Anticoagulant — bleeding risk'),
    ('DOPAMINE', 'µg/kg/min', '3.2 mg/ml', '2 / 20', '1 / 50', 'Vasopressor/inotrope — dose-dependent receptor activity'),
    ('NORADRENALINE', 'µg/kg/min', '0.08 mg/ml', '0.01 / 0.3', '0.001 / 2.0', 'Vasopressor — septic shock, extreme narrow window'),
    ('PROPOFOL', 'mg/kg/h', '10 mg/ml', '1 / 6', '0.5 / 12', 'Sedative/anaesthetic — PRIS risk above 4 mg/kg/h'),
    ('INSULIN', 'U/h', '1 U/ml', '1 / 10', '0.5 / 50', 'Hypoglycaemia risk — requires glucose monitoring'),
    ('AMIODARONE', 'mg/h', '1.8 mg/ml', '10 / 100', '5 / 150', 'Antiarrhythmic — phlebitis risk'),
    ('KCl 20mmol', 'mmol/h', '1 mmol/ml', '5 / 20', '1 / 40', 'Electrolyte — cardiac arrest risk if rapid infusion'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('9.2 B. Braun Drug Library (braunDrugLibrary.ts)')
body(
    'The B. Braun drug library mirrors the Alaris GP library in drug selection but is adapted '
    'to the B. Braun unit system and clinical profiles. Notable differences: '
    'advisory limits (pre-warning tier) added below soft limits for each drug.'
)

doc.add_paragraph()
h2('9.3 Graseby 3100 — No Drug Library')
body(
    'The Graseby 3100 has no drug library. The sessionAdapter.ts maps it to the "manual" drug '
    'entry from the Alaris GP library (drug_library_used = 0 always). This correctly sets '
    'guardrail_soft_min/max/hard_min/max to the full range [0.1, 199.9], '
    'and ensures R12 (no library) always fires, placing every session at minimum MEDIUM risk.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. MANUAL REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
h1('10. Manual References & Source Documents')

h2('10.1 Documents Included in docs/ Directory')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
table_header(t, ['Filename', 'Device', 'Type', 'Extractable Text?'])
for row in [
    ('alaris-infusion-central-user-manual---eng.pdf', 'Alaris GP / Infusion Central', 'User Manual', 'Yes'),
    ('BD_Alaris_8015_PC_User_Manual_2020-01.pdf', 'BD Alaris 8015 PC', 'User Manual', 'Yes'),
    ('braun_infusomat_space_ifu_686N_GB.pdf', 'B. Braun Infusomat Space (GB)', 'IFU', 'Yes'),
    ('braun_infusomat_space_ifu_586U_US.pdf', 'B. Braun Infusomat Space (US)', 'IFU', 'Yes'),
    ('braun_infusomat_service_manual.pdf', 'B. Braun Infusomat Space', 'Service Manual', 'Yes'),
    ('braun_infusomat_specs.md', 'B. Braun Infusomat Space', 'Extracted Spec Notes', 'N/A (Markdown)'),
    ('graseby_3100_technical_service_manual_2004.pdf', 'Graseby 3100', 'Technical Service Manual (00SM-0131-7, Issue 7)', 'Yes'),
    ('graseby_3300_pca_technical_service_manual_2004.pdf', 'Graseby 3300 PCA', 'Technical Service Manual (00SM-0128-7, Issue 7)', 'Yes'),
]:
    add_row(t, row, shade=row[2] == 'Technical Service Manual (00SM-0131-7, Issue 7)')

doc.add_paragraph()
h2('10.2 Design Decision Traceability')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Design Decision', 'Source Document', 'Location'])
for row in [
    ('RATE_MIN=0.1, RATE_MAX=1200 (Alaris)', 'BD DFU 1000DF00152 Issue 1', 'Factory Default Data Set table'),
    ('VTBI_MAX=9999, KVO_RATE=1.0', 'BD DFU 1000DF00152 Issue 1', 'Factory Default Data Set table'),
    ('BOLUS_RATE_DEFAULT=500, BOLUS_VOLUME_MAX=5', 'BD DFU BDDF00535 Issue 4', 'Guardrails specification'),
    ('STEP_LARGE=10, STEP_SMALL=1', 'BD DFU 1000DF00152 Issue 1', 'Section 3 — faster/slower description'),
    ('HOLD_DELAY_MS=500, HOLD_REPEAT_MS=80', 'BD DFU 1000DF00152 Issue 1', 'Section 3 — hold-to-accelerate'),
    ('PRESSURE_DEFAULT=5 (L5)', 'BD DFU 1000DF00152 Issue 1', 'Factory Default Data Set — Pressure Default L5'),
    ('AIL_LIMIT_MAX=100µl', 'BD DFU 1000DF00152 Issue 1', 'Factory Default Data Set — AIL Limit Max 100µl'),
    ('Graseby RATE_MAX=199.9', 'Graseby 3100 Technical Service Manual 00SM-0131-7', 'Chapter 1 Specification — Flow rate 0.1 to 199.9 ml/h'),
    ('Graseby SyringeCapacityMl = 20|30|50', 'Graseby 3100 Technical Service Manual 00SM-0131-7', 'Chapter 1 Specification — Useable syringes: 20, 30, 50/60 ml'),
    ('Guardrail soft/hard limits (Alaris drug library)', 'BD DFU BDDF00535 Issue 4', 'Guardrails — per-drug limit tables'),
    ('B. Braun AIL sensitivity 20–500µl', 'B. Braun Infusomat Space IFU 686N GB', 'Technical specifications section'),
    ('CVE-2021-33885 CVSS 9.0 (B. Braun)', 'CISA ICSMA-21-294-01', 'Published NVD/CISA advisory'),
    ('CVE-2022-22772 CVSS 9.8 (BD Alaris)', 'CISA ICSMA-23-194-01', 'Published NVD/CISA advisory'),
    ('golden_path_ratio definition', 'Cauchi et al. (2011) EICS4Med', 'Section 3 — Minimum keypress model'),
    ('relative_error threshold (Thimbleby)', 'Thimbleby & Cairns (2010) JRSI 7(51)', 'Section 4 — Dose error magnitude analysis'),
    ('NPSA 2010 — Graseby 3100 clinical risk', 'NPSA Patient Safety Alert (2010, UK)', 'Alert NPSA/2010/RRR019'),
]:
    add_row(t, row)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. RISK RULES
# ═══════════════════════════════════════════════════════════════════════════════
h1('11. Risk Rule Reference (R01–R21)')

body(
    'All labelling rules are implemented in src/ai/labellingRules.ts. '
    'Rules are applied to each TrainingRecord after feature extraction. '
    'HIGH rules (any single trigger → "high") take priority over MEDIUM rules.'
)

doc.add_paragraph()
h2('11.1 HIGH Risk Rules (R01–R08)')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
table_header(t, ['Rule ID', 'Label', 'Test Condition', 'Score Weight'])
for row in [
    ('R01', '10x dose error', 'relative_error >= 0.9', '1.00'),
    ('R02', '50% dose error', 'relative_error >= 0.5', '0.85'),
    ('R03', 'Hard guardrail limit triggered', 'guardrail_blocked === 1', '0.80'),
    ('R04', 'MANUAL mode + high rate (>200 ml/h)', 'drug_library_used===0 AND final_rate>200', '0.75'),
    ('R05', 'Firmware version has known CVE', 'firmware_version_risk === 1', '0.70'),
    ('R06', 'No maintenance in >1 year', 'days_since_maintenance > 365', '0.65'),
    ('R07', 'KCl above soft limits', 'drug_id==="kcl" AND rate_within_soft_limits===0', '0.90'),
    ('R08', 'Confirmed wrong value started', 'confirmed_incorrect===1 AND relative_error>0.1', '0.80'),
]:
    add_row(t, row, shade=True)

doc.add_paragraph()
h2('11.2 MEDIUM Risk Rules (R10–R21)')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
table_header(t, ['Rule ID', 'Label', 'Test Condition', 'Score Weight'])
for row in [
    ('R10', '10% dose error', 'relative_error >= 0.1', '0.40'),
    ('R11', 'Soft guardrail warning overridden', 'guardrail_override === 1', '0.45'),
    ('R12', 'MANUAL mode — no library protection', 'drug_library_used === 0', '0.30'),
    ('R13', 'VTBI not set', 'vtbi_set === 0', '0.25'),
    ('R14', 'Excessive corrections (>5 reversals)', 'correction_count > 5', '0.30'),
    ('R15', 'Boundary hit during rate entry', 'boundary_hit_count > 0', '0.20'),
    ('R16', 'Drug library outdated (>90 days)', 'drug_library_age_days > 90', '0.25'),
    ('R17', 'Battery critically low (<20%)', 'battery_level_pct < 20', '0.35'),
    ('R18', 'Not connected to Gateway', 'network_connected === 0', '0.30'),
    ('R19', 'High configuration drift (>0.5)', 'config_drift_score > 0.5', '0.30'),
    ('R20', 'Recent occlusion alarms (>=3)', 'recent_occlusion_alarms >= 3', '0.40'),
    ('R21', 'Slow rate entry (>60 seconds)', 'entry_time_ms > 60000', '0.20'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('11.3 Layer Mapping')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Layer', 'Rule IDs', 'Clinical Basis'])
for row in [
    ('Layer 1 — Interaction', 'R01, R02, R08, R10, R14, R15, R21', 'Cauchi et al. (2011); Thimbleby & Cairns (2010)'),
    ('Layer 2 — Configuration', 'R03, R04, R07, R11, R12, R13', 'ISMP High-Alert Medications; IEC 60601-2-24:2012'),
    ('Layer 3 — System', 'R05, R06, R16, R17, R18, R19, R20', 'FDA Infusion Pump Safety; CISA ICS-CERT advisories'),
]:
    add_row(t, row)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. DEVICE DESIGN REGISTRY
# ═══════════════════════════════════════════════════════════════════════════════
h1('12. Device Design Score Registry (deviceDesign.ts)')

body(
    'The registry stores static design-time properties for 7 infusion pump models. '
    '3 are fully simulated (Alaris GP, B. Braun, Graseby 3100). '
    '4 are stubs with design profiles for Phase 2 and 3 expansion (Sigma Spectrum, Plum 360, Agilia VP, CADD-Solis).'
)
t = doc.add_table(rows=1, cols=8)
t.style = 'Table Grid'
table_header(t, ['Device', 'Drug Lib', 'Guard Tiers', 'VTBI', 'Network', 'CVEs', 'Max CVSS', 'Status'])
for row in [
    ('BD Alaris GP', 'Yes', '2', 'Yes', 'Yes', '13', '9.8', 'SIMULATED'),
    ('B. Braun Infusomat Space', 'Yes', '3', 'Yes', 'Yes', '5', '9.0', 'SIMULATED'),
    ('Graseby 3100', 'No', '0', 'No', 'No', '0', '0', 'SIMULATED'),
    ('Baxter Sigma Spectrum', 'Yes', '2', 'Yes', 'Yes', '4', '7.5', 'Stub — Phase 2'),
    ('ICU Medical Plum 360', 'Yes', '2', 'Yes', 'Yes', '0', '0', 'Stub — Phase 2'),
    ('Fresenius Agilia VP', 'Yes', '2', 'Yes', 'Yes', '13', '7.5', 'Stub — Phase 3'),
    ('CADD-Solis PCA', 'Yes', '2', 'Yes', 'Yes', '2', '9.9', 'Stub — Phase 3'),
]:
    add_row(t, row, shade=row[7] == 'SIMULATED')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  13. TRAINING RECORD SCHEMA
# ═══════════════════════════════════════════════════════════════════════════════
h1('13. Feature Schema — TrainingRecord')

body(
    'Every simulator session produces a TrainingRecord with 55 fields across 5 sections. '
    'All fields are serialisable to JSON and CSV without transformation. '
    'Implemented in src/ai/featureExtractor.ts.'
)

h2('13.1 Metadata (5 fields)')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Field', 'Type', 'Description'])
for row in [
    ('record_id', 'string', 'Unique record identifier (e.g. "human_alaris_gp_1714800000000")'),
    ('session_id', 'string', 'UUID v4 from Web Crypto API'),
    ('timestamp_iso', 'string', 'ISO 8601 datetime of session capture'),
    ('pump_model', 'string', '"alaris_gp" | "braun_infusomat" | "graseby_3100"'),
    ('firmware_version', 'string', 'From device constants or scenario generator'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('13.2 Interaction Features (14 fields) — Layer 1')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Field', 'Type', 'Description'])
for row in [
    ('entry_time_ms', 'number', 'ms from first keypress to infusion_started'),
    ('total_keypresses', 'number', 'All rate_adjust events during entry'),
    ('large_up_count', 'number', '»» presses (delta === +10)'),
    ('small_up_count', 'number', '» presses (delta === +1)'),
    ('small_down_count', 'number', '« presses (delta === -1)'),
    ('large_down_count', 'number', '«« presses (delta === -10)'),
    ('correction_count', 'number', 'Direction reversal events in log'),
    ('boundary_hit_count', 'number', 'Times RATE_MIN or RATE_MAX was hit'),
    ('large_btn_ratio', 'number', '(large_up + large_down) / total_keypresses'),
    ('golden_path_ratio', 'number', 'actual / minimum keypresses (Cauchi et al. 2011)'),
    ('final_rate_ml_h', 'number', 'Programmed rate converted to ml/h'),
    ('intended_rate_ml_h', 'number', 'Target rate from task mode or scenario'),
    ('error_magnitude_ml_h', 'number', 'abs(final - intended) in ml/h'),
    ('relative_error', 'number', 'error / intended (capped 2.0) — Thimbleby & Cairns'),
    ('confirmed_incorrect', '0 | 1', '1 if infusion started with >25% error'),
    ('drug_unit_used', 'string', 'The unit displayed during entry (e.g. "µg/kg/min")'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('13.3 Configuration Features (16 fields) — Layer 2')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Field', 'Type', 'Description'])
for row in [
    ('drug_id', 'string', 'drug.id from DRUG_LIBRARY'),
    ('drug_name', 'string', 'drug.name (max 12 chars)'),
    ('drug_library_used', '0 | 1', '1 if drug selected from library, 0 if MANUAL'),
    ('guardrail_soft_min/max', 'number', 'Drug soft limits in drug units'),
    ('guardrail_hard_min/max', 'number', 'Drug hard limits in drug units'),
    ('guardrail_warning_shown', '0 | 1', '1 if soft limit warning appeared'),
    ('guardrail_override', '0 | 1', '1 if OVERRIDE chosen — safety-critical event'),
    ('guardrail_blocked', '0 | 1', '1 if hard limit triggered'),
    ('rate_within_soft_limits', '0 | 1', '1 if final rate within soft limits'),
    ('vtbi_set', '0 | 1', '1 if VTBI was programmed before RUN'),
    ('vtbi_value_ml', 'number', 'VTBI in ml (0 if not set)'),
    ('kvo_rate_ml_h', 'number', 'KVO rate configured (1.0 ml/h default)'),
    ('bolus_delivered', '0 | 1', '1 if bolus was used during session'),
    ('bolus_volume_ml', 'number', 'Total bolus volume in ml'),
    ('patient_weight_kg', 'number', 'Weight for dose/kg calculation'),
    ('pressure_alarm_level', 'number', 'L0–L8, default L5'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('13.4 System Features (8 fields) — Layer 3')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Field', 'Type', 'Description'])
for row in [
    ('days_since_maintenance', 'number', '0–730 (2 years max)'),
    ('battery_level_pct', 'number', '0–100%'),
    ('firmware_version_risk', '0 | 1', '1 if firmware in known CVE list'),
    ('network_connected', '0 | 1', '1 if connected to Gateway'),
    ('drug_library_age_days', 'number', 'Days since drug library last updated'),
    ('config_drift_score', 'number', '0.0–1.0 distance from hospital standard config'),
    ('recent_occlusion_alarms', 'number', 'Count in simulated last 24h'),
    ('alarms_during_session', 'number', 'Non-INFUSION_COMPLETE alarms in this session'),
]:
    add_row(t, row)

doc.add_paragraph()
h2('13.5 Risk & Score Fields (13 fields)')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
table_header(t, ['Field', 'Type', 'Description'])
for row in [
    ('risk_label', '"low"|"medium"|"high"', 'Derived from grade'),
    ('risk_score', 'number', 'Equals composite_score (0–1)'),
    ('risk_reasons', 'string[]', 'All R-rule IDs that fired (pipe-separated in CSV)'),
    ('design_score', 'number', 'Layer 0 (0–1) — fixed per device model'),
    ('interaction_score', 'number', 'Layer 1 (0–1) — programming behaviour'),
    ('configuration_score', 'number', 'Layer 2 (0–1) — session setup'),
    ('system_score', 'number', 'Layer 3 (0–1) — device physical/cyber state'),
    ('composite_score', 'number', '0.20×L0 + 0.30×L1 + 0.25×L2 + 0.25×L3'),
    ('grade', 'A+|A|B|C|D|E|F', 'Energy-label style grade from composite_score'),
    ('design_reasons', 'string[]', 'Design factor penalties that fired'),
    ('interaction_reasons', 'string[]', 'Layer 1 R-rules that fired'),
    ('configuration_reasons', 'string[]', 'Layer 2 R-rules that fired'),
    ('system_reasons', 'string[]', 'Layer 3 R-rules that fired'),
]:
    add_row(t, row)

page_break()

# ── Footer on last page ────────────────────────────────────────────────────────
hr()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    f'PumpSimulator2 — Technical Documentation  |  '
    f'Generated {datetime.date.today().strftime("%d %B %Y")}  |  '
    f'QR Seed Pilot Study — University Research'
)
run.font.size  = Pt(8)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = '/home/nandhakumar/Documents/PumpSimulator2/docs/PumpSimulator2_Documentation.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
